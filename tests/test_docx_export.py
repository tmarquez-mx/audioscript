import os
import unittest
from unittest.mock import patch

import pandas as pd
from docx import Document

import App


class DocxExportTests(unittest.TestCase):
    def test_exports_transcript_with_analysis_and_formatting(self):
        memos = pd.DataFrame([{"segmento": 1, "memo": "Lectura analítica"}])
        codes = pd.DataFrame(
            [
                {
                    "segmento": 1,
                    "codigo": "identidad",
                    "cita": "frase seleccionada",
                    "nota": "Nota del código",
                }
            ]
        )

        path = App.save_as_docx(
            "Entrevistadora: Texto **negrita**, __subrayado__, ==resaltado== y ~~tachado~~.",
            "Entrevista de prueba",
            "30/06/2026",
            "30/06/2026 20:00",
            "Contexto de prueba",
            memos,
            codes,
        )
        self.addCleanup(lambda: os.path.exists(path) and os.unlink(path))

        document = Document(path)
        document_text = "\n".join(p.text for p in document.paragraphs)
        self.assertIn("Entrevista de prueba", document_text)
        self.assertIn("Lectura analítica", document_text)
        self.assertIn("identidad", document_text)

    def test_footer_template_failure_does_not_block_export(self):
        with patch.object(App, "add_docx_footer", side_effect=FileNotFoundError("template")):
            path = App.save_as_docx(
                "Texto recuperable",
                "Exportación resiliente",
                "",
                "",
                "",
                pd.DataFrame(),
                pd.DataFrame(),
            )
        self.addCleanup(lambda: os.path.exists(path) and os.unlink(path))

        document = Document(path)
        self.assertIn("Texto recuperable", "\n".join(p.text for p in document.paragraphs))


if __name__ == "__main__":
    unittest.main()
