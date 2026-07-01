import os
import re
import json
import html
import math
import io
import shutil
import tempfile
import hashlib
import uuid
from datetime import datetime, timedelta
import xml.etree.ElementTree as ET
import importlib.util
import logging

import pandas as pd
import streamlit as st
from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

from styles import apply_global_styles
from shared_config import (
    BYTES_PER_MB,
    DEFAULT_TRANSCRIPTION_LANGUAGE,
    DEFAULT_WHISPER_MODEL,
    GUIDED_INSTALL_MODEL_OPTIONS,
    MODE_COMPLETE,
    MODE_SEGMENTED,
    PROCESSING_MODES,
    TRANSCRIPTION_LANGUAGES,
    normalize_processing_mode,
    sanitize_filename,
)
from transcribe import (
    build_initial_prompt,
    configure_ssl_certificates,
    ensure_whisper_model_installed,
    get_app_support_dir,
    get_whisper_model_installation_info,
    is_whisper_model_installed,
    load_whisper_model,
    transcribe_audio,
)
from media import (
    check_ffmpeg,
    configure_ffmpeg,
    detect_silence_ranges,
    estimate_quick_split_preview,
    format_duration,
    get_media_metadata,
    split_audio,
    split_large_media_by_size,
)
from data_model import (
    build_material_id, build_named_project_id, build_project_id,
    build_project_settings_from_session, delete_material,
    delete_project_state, display_media_name, empty_material_transcription_state,
    format_size_mb, get_active_material_record, get_default_project_settings,
    get_default_session_setting_state,
    get_material_dir, get_material_record, get_material_transcription_path,
    get_project_audio_path, get_project_chunks_dir, get_project_dir,
    get_project_legacy_state_path, get_project_materials_root_dir,
    get_project_media_dir, get_project_media_path, get_project_meta_path,
    get_project_quick_split_dir, get_project_state_path, get_relative_project_path,
    list_existing_chunks, list_project_materials, list_project_media,
    load_material_transcription, load_project_meta, load_project_state,
    merge_v2_project_state, migrate_project_v1_to_v2,
    normalize_project_date, now_iso, rename_project, resolve_project_path,
    project_settings_to_session_state,
    save_legacy_state_document,
    save_material_transcription_document, save_project_meta_document,
    scan_project_media_files, update_material_progress,
)
from ui_components import (
    apply_editor_styles, close_analysis_desk, close_card, highlight_transcribe_buttons,
    keep_sidebar_accessible, open_analysis_desk, open_card,
        render_ai_margin_panel, render_app_header, render_audio_console_header,
        render_beta_installer_preview,
        render_code_selection_from_editor, render_config_separator, render_footer_bar,
        render_dom_script,
        render_font_size_dial, render_keyboard_shortcuts, render_rotary_dial,
    render_segment_audio_player, render_segment_minutes_dial,
    render_segmented_or_fallback, render_selectable_transcript_panel,
    render_workspace_watermarks,
    render_sidebar_circle, render_sidebar_minute_ring, render_transcript_panel_label,
    render_transcript_toolbar_title, select_dial_option,
)
try:
    import certifi
except ImportError:
    certifi = None


st.set_page_config(
    page_title="AudioScript Contextual",
    layout="wide",
    initial_sidebar_state="expanded",
)

apply_global_styles()



TEMP_DIR = os.path.join(tempfile.gettempdir(), "audioscript_contextual")
os.makedirs(TEMP_DIR, exist_ok=True)
QUICK_SPLIT_ROOT = os.path.join(tempfile.gettempdir(), "audioscript_splits")
os.makedirs(QUICK_SPLIT_ROOT, exist_ok=True)
LOGGER = logging.getLogger("audioscript")
WHISPER_FILE_LIMIT_MB = 200
WHISPER_FILE_LIMIT_BYTES = WHISPER_FILE_LIMIT_MB * BYTES_PER_MB
APP_DIR = os.path.dirname(os.path.abspath(__file__))
INSTALLER_COMPLETED_MARKER = os.path.join(get_app_support_dir(), ".installer-onboarding-complete")
PROJECTS_DIR = os.environ.get(
    "AUDIOSCRIPT_DATA_DIR",
    os.path.join(APP_DIR, ".audioscript_projects"),
)
os.makedirs(PROJECTS_DIR, exist_ok=True)
@st.cache_resource(show_spinner=False)
def check_ffmpeg_once():
    """Evita relanzar ffmpeg -version en cada rerun de la interfaz."""
    return check_ffmpeg()








def render_project_context_card():
    """Muestra el contexto activo como ficha de mesa de trabajo."""
    ensure_transcription_stamp()
    st.markdown(
        f"""
        <div class="section-card section-card--transcription">
          <div class="section-title">Mesa del proyecto</div>
          <div style="display:grid;grid-template-columns:repeat(5,minmax(0,1fr));gap:.75rem;">
            <div><div class="transcript-toolbar-label">Proyecto</div><strong>{html.escape(st.session_state.main_doc_title)}</strong></div>
            <div><div class="transcript-toolbar-label">Audio/video activo</div><strong>{html.escape(st.session_state.active_audio_name or "Material guardado")}</strong></div>
            <div><div class="transcript-toolbar-label">Fecha del evento</div><strong>{st.session_state.main_event_date.strftime("%d/%m/%Y") if st.session_state.main_event_date else "No especificada"}</strong></div>
            <div><div class="transcript-toolbar-label">Fecha de transcripción</div><strong>{html.escape(st.session_state.main_transcription_date or "No registrada")}</strong></div>
            <div><div class="transcript-toolbar-label">Segmentos</div><strong>{len(st.session_state.transcript_segments)}</strong></div>
          </div>
        </div>
        """,
        unsafe_allow_html=True,
    )


































































































def build_project_meta_from_session(project_id, existing_meta=None):
    existing_meta = existing_meta or {}
    created_value = existing_meta.get("created", now_iso())
    settings = build_project_settings_from_session(
        st.session_state,
        existing_meta.get("settings", get_default_project_settings()),
    )
    return {
        "schema_version": 2,
        "project_id": project_id,
        "title": st.session_state.main_doc_title,
        "description": st.session_state.get("project_description", ""),
        "event_date": normalize_project_date(st.session_state.main_event_date),
        "created": created_value,
        "updated": now_iso(),
        "active_material_id": st.session_state.get("active_material_id", ""),
        "settings": settings,
        "materials": existing_meta.get("materials", []),
        "quick_split_last_summary": st.session_state.get("quick_split_last_summary", {}),
        "quick_session_mode": st.session_state.get("quick_session_mode", False),
        "main_transcription_date": st.session_state.main_transcription_date,
        "autosave_last_saved": datetime.now().strftime("%d/%m/%Y %H:%M:%S"),
    }


def apply_material_transcription_state(material_state):
    material_state = material_state or empty_material_transcription_state()
    st.session_state.current_chunk_idx = material_state.get("current_chunk_idx", 0)
    st.session_state.transcript_segments = material_state.get("transcript_segments", [])
    st.session_state.segment_texts = material_state.get("segment_texts", [])
    st.session_state.current_segment_text = material_state.get("current_segment_text", "")
    if (
        st.session_state.segment_texts
        and 0 <= st.session_state.current_chunk_idx < len(st.session_state.segment_texts)
    ):
        st.session_state.current_segment_text = st.session_state.segment_texts[
            st.session_state.current_chunk_idx
        ]
    st.session_state.memos = material_state.get("memos", [])
    st.session_state.codes = material_state.get("codes", [])
    st.session_state.last_transcription = material_state.get("last_transcription", "")
    st.session_state.chunks_prepared = material_state.get("chunks_prepared", False)
    st.session_state.chunk_segment_mins = material_state.get("chunk_segment_mins")




def material_transcription_from_session(material_id):
    if st.session_state.get("chunks"):
        sync_current_segment_text()
    return {
        "material_id": material_id,
        "chunks_prepared": st.session_state.chunks_prepared,
        "chunk_segment_mins": st.session_state.chunk_segment_mins,
        "current_chunk_idx": st.session_state.current_chunk_idx,
        "transcript_segments": st.session_state.transcript_segments,
        "segment_texts": st.session_state.get("segment_texts", []),
        "current_segment_text": st.session_state.current_segment_text,
        "memos": st.session_state.memos,
        "codes": st.session_state.codes,
        "last_transcription": st.session_state.last_transcription,
        "updated": now_iso(),
    }


def save_material_transcription(project_id, material_id):
    if not project_id or not material_id:
        return
    save_material_transcription_document(
        project_id,
        material_id,
        material_transcription_from_session(material_id),
    )


def save_project_meta(project_id):
    if not project_id:
        return None
    os.makedirs(get_project_dir(project_id), exist_ok=True)
    project_meta = load_project_meta(project_id) or build_project_meta_from_session(project_id)
    project_meta = build_project_meta_from_session(project_id, project_meta)
    active_material_id = st.session_state.get("active_material_id")
    active_material = get_material_record(project_meta, active_material_id) if active_material_id else None
    if active_material:
        update_material_progress(active_material)
        active_material["updated"] = now_iso()
    save_project_meta_document(project_id, project_meta)
    st.session_state.autosave_last_saved = project_meta["autosave_last_saved"]
    return project_meta




















def render_quick_split_wave_preview(segments, duration_seconds):
    """Dibuja una onda sintetica con marcas de corte para el preview."""
    bar_count = 48
    bars = []
    for index in range(bar_count):
        height = 22 + int(24 * abs(math.sin(index * 0.52)) + 10 * abs(math.sin(index * 1.13)))
        x_position = 8 + index * 8
        y_position = 42 - height / 2
        bars.append(
            f'<rect x="{x_position}" y="{y_position:.1f}" width="4" height="{height}" rx="2" fill="rgba(255,255,255,0.56)" />'
        )

    cut_lines = []
    for index, segment in enumerate(segments[:-1], start=1):
        cut_at = segment.get("end")
        if duration_seconds and duration_seconds > 0 and cut_at:
            x_position = 8 + (cut_at / duration_seconds) * 376
        else:
            x_position = 8 + (index / len(segments)) * 376
        cut_lines.append(
            f'<line x1="{x_position:.1f}" y1="8" x2="{x_position:.1f}" y2="76" stroke="#38b795" stroke-width="2.5" stroke-linecap="round" />'
        )

    return (
        '<div class="quick-split-wave" aria-label="Previsualización de onda y cortes">'
        '<svg viewBox="0 0 400 84" width="100%" height="84" role="img">'
        '<rect x="0" y="0" width="400" height="84" rx="12" fill="#12395f"></rect>'
        f'{"".join(bars)}'
        f'{"".join(cut_lines)}'
        '</svg>'
        '</div>'
    )


def render_quick_split_preview(source_label, source_size, max_mb, use_silence, key_prefix):
    """Muestra metadatos, onda y chips de segmentos estimados antes de dividir."""
    metadata = {}
    active_path = st.session_state.active_audio_path
    if key_prefix == "active_material" and active_path and os.path.exists(active_path):
        metadata = get_media_metadata(active_path)

    duration_seconds = metadata.get("duration_seconds")
    segments = estimate_quick_split_preview(source_size, int(max_mb), duration_seconds)
    average_size = sum(segment["size_bytes"] for segment in segments) / len(segments)
    average_duration = (
        (duration_seconds / len(segments))
        if duration_seconds and segments
        else None
    )
    chips = []
    for segment in segments[:8]:
        time_range = (
            f"{format_duration(segment['start'])}-{format_duration(segment['end'])}"
            if segment.get("start") is not None and segment.get("end") is not None
            else "rango estimado"
        )
        chips.append(
            f'<span class="quick-split-chip">Parte {segment["index"]}: {format_size_mb(segment["size_bytes"])} · {time_range}</span>'
        )
    if len(segments) > 8:
        chips.append(f'<span class="quick-split-chip">+{len(segments) - 8} más</span>')

    silence_text = (
        "Los cortes buscarán pausas naturales al ejecutar la división."
        if use_silence
        else "La división usará cortes regulares sin búsqueda de pausas."
    )
    route_preview = f"/tmp/audioscript_splits/{sanitize_filename(os.path.splitext(source_label)[0])}_parteXXX.mp3"
    wave_preview = render_quick_split_wave_preview(segments, duration_seconds)
    chips_html = "".join(chips)
    preview_html = (
        '<div class="quick-split-preview">'
        '<div class="quick-split-meta">'
        f'<span><strong>Duración</strong><br>{format_duration(duration_seconds)}</span>'
        f'<span><strong>Tamaño</strong><br>{format_size_mb(source_size)}</span>'
        f'<span><strong>Bitrate</strong><br>{html.escape(metadata.get("bitrate") or "pendiente")}</span>'
        f'<span><strong>Frecuencia</strong><br>{html.escape(metadata.get("sample_rate") or "pendiente")}</span>'
        '</div>'
        f'{wave_preview}'
        f'<div class="quick-split-chip-row">{chips_html}</div>'
        '<div class="quick-split-summary">'
        f'<strong>Preview:</strong> {len(segments)} segmento(s), tamaño promedio {format_size_mb(average_size)}, duración promedio {format_duration(average_duration)}.<br>'
        f'{silence_text}<br>'
        f'<strong>Ruta temporal:</strong> {html.escape(route_preview)}'
        '</div>'
        '</div>'
    )
    st.markdown(preview_html, unsafe_allow_html=True)


def render_dialog_split_recommendation(source_size, max_mb, use_silence):
    """Dibuja en el dialogo una vista compacta, tipo maqueta, de los segmentos sugeridos."""
    segments = estimate_quick_split_preview(source_size, int(max_mb))
    largest_segment = max((segment["size_bytes"] for segment in segments), default=1)
    bars = []
    for segment in segments:
        width_percent = max(12, int((segment["size_bytes"] / largest_segment) * 100))
        bars.append(
            f'<i title="Parte {segment["index"]}: {format_size_mb(segment["size_bytes"])}" style="flex: {width_percent} 1 0;"></i>'
        )

    silence_label = "silencios activados" if use_silence else "cortes regulares"
    average_size = sum(segment["size_bytes"] for segment in segments) / len(segments)
    bars_html = "".join(bars)
    st.markdown(
        f"""
        <div class="dialog-split-card">
          <strong>Conviene dividir este archivo</strong>
          <span>Supera el límite recomendado. AudioScript lo parte en segmentos y busca silencios para no cortar a media palabra.</span>
          <div class="dialog-split-stripes" aria-label="Representación visual de segmentos sugeridos">
            {bars_html}
          </div>
          <div class="dialog-split-caption">
            Cada bloque representa un segmento estimado. Los colores alternos ayudan a distinguir los cortes antes de activar la división.
          </div>
          <div class="dialog-split-summary-row">
            <span class="dialog-split-chip">{len(segments)} segmentos</span>
            <span class="dialog-split-chip">promedio {format_size_mb(average_size)}</span>
            <span class="dialog-split-chip">{silence_label}</span>
          </div>
        </div>
        """,
        unsafe_allow_html=True,
    )


def get_dialog_quick_split_settings(key_prefix):
    """Lee los ajustes compactos del dialogo antes de dibujarlos."""
    choice_key = f"{key_prefix}_quick_split_size_choice"
    custom_key = f"{key_prefix}_quick_split_custom_size"
    silence_key = f"{key_prefix}_quick_split_use_silence"
    size_choice = st.session_state.get(choice_key, WHISPER_FILE_LIMIT_MB)
    if size_choice == "Personalizado":
        max_mb = int(st.session_state.get(custom_key, st.session_state.quick_split_max_mb))
    else:
        try:
            max_mb = int(size_choice)
        except (TypeError, ValueError):
            max_mb = int(st.session_state.quick_split_max_mb)
    use_silence = bool(st.session_state.get(silence_key, True))
    return max_mb, use_silence


def render_dialog_split_settings(source_label, source_size, key_prefix):
    """Ajustes compactos de division para no saturar el dialogo de material."""
    st.caption(
        f"{source_label} pesa {format_size_mb(source_size)}. Puedes ajustar el tamaño máximo de cada parte antes de activar."
    )
    size_choice = st.radio(
        "Tamaño máximo por segmento",
        [25, 50, 100, 200, 500, "Personalizado"],
        index=3,
        horizontal=True,
        key=f"{key_prefix}_quick_split_size_choice",
    )
    if size_choice == "Personalizado":
        max_mb = st.number_input(
            "Tamaño personalizado en MB",
            min_value=10,
            max_value=900,
            value=int(st.session_state.quick_split_max_mb),
            step=10,
            key=f"{key_prefix}_quick_split_custom_size",
        )
    else:
        max_mb = int(size_choice)
    st.session_state.quick_split_max_mb = int(max_mb)
    use_silence = st.checkbox(
        "Buscar silencio antes del corte",
        value=True,
        key=f"{key_prefix}_quick_split_use_silence",
        help="Evita cortar palabras a la mitad detectando pausas naturales cerca del punto de corte.",
    )
    return int(max_mb), use_silence




def mark_material_needs_preparation(project_meta):
    """Al cambiar material, la preparación debe confirmarse de nuevo."""
    settings = project_meta.setdefault("settings", get_default_project_settings())
    settings["prep_confirmed"] = False
    settings["transcription_started"] = False
    st.session_state.expand_preparation_panel = True
    st.session_state.scroll_to_preparation_panel = True


def activate_existing_material(material_ref):
    """Cambia el material activo del proyecto sin borrar transcripciones previas."""
    project_id = st.session_state.current_project_id
    if not project_id:
        return

    project_meta = load_project_meta(project_id)
    if not project_meta or project_meta.get("schema_version") != 2:
        project_state = load_project_state(project_id)
        project_meta = load_project_meta(project_id)
        if not project_meta:
            return

    current_material_id = st.session_state.get("active_material_id")
    if current_material_id:
        save_material_transcription(project_id, current_material_id)
        save_project_meta(project_id)

    if os.path.exists(str(material_ref)):
        material_record = next(
            (
                item for item in project_meta.get("materials", [])
                if resolve_project_path(project_id, item.get("path", "")) == material_ref
            ),
            None,
        )
    else:
        material_record = get_material_record(project_meta, str(material_ref))

    if not material_record:
        return

    reset_transcription_state()
    st.session_state.active_material_id = material_record.get("id")
    st.session_state.active_audio_name = material_record.get("name", "")
    st.session_state.active_audio_path = resolve_project_path(project_id, material_record.get("path", ""))
    st.session_state.uploaded_file_id = material_record.get("id")
    st.session_state.show_new_project_uploader = False
    apply_material_transcription_state(
        load_material_transcription(project_id, material_record.get("id"))
    )
    project_meta["active_material_id"] = material_record.get("id")
    mark_material_needs_preparation(project_meta)
    save_project_meta(project_id)
    save_project_state()


def get_project_audio_source(project_id):
    """Encuentra el audio guardado dentro de un proyecto."""
    project_dir = get_project_dir(project_id)
    if not os.path.isdir(project_dir):
        return ""

    for filename in sorted(os.listdir(project_dir)):
        if filename.startswith("source"):
            return os.path.join(project_dir, filename)
    return ""


def activate_uploaded_material(uploaded_file):
    """Agrega o cambia el audio/video activo dentro del proyecto actual."""
    project_id = st.session_state.current_project_id
    if not project_id:
        st.warning("Primero crea o abre un proyecto; después agrega el audio o video.")
        return None

    project_meta = load_project_meta(project_id) or build_project_meta_from_session(project_id)
    if st.session_state.get("active_material_id"):
        save_material_transcription(project_id, st.session_state.active_material_id)

    destination_path = get_project_media_path(project_id, uploaded_file.name)
    save_uploaded_file(uploaded_file, destination_path)
    metadata = get_media_metadata(destination_path)
    material_id = build_material_id(
        uploaded_file.name,
        f"{uploaded_file.name}-{uploaded_file.size}-{now_iso()}",
    )
    material_record = {
        "id": material_id,
        "name": uploaded_file.name,
        "path": get_relative_project_path(project_id, destination_path),
        "size_bytes": metadata.get("size_bytes", uploaded_file.size),
        "duration_seconds": metadata.get("duration_seconds"),
        "added": now_iso(),
        "source": "upload",
        "split_summary": None,
        "total_segments": 0,
        "transcribed_segments": 0,
        "status": "pendiente",
        "updated": now_iso(),
    }
    project_meta.setdefault("materials", []).append(material_record)
    project_meta["active_material_id"] = material_id
    mark_material_needs_preparation(project_meta)

    reset_transcription_state()
    st.session_state.active_material_id = material_id
    st.session_state.uploaded_file_id = material_id
    st.session_state.active_audio_name = uploaded_file.name
    st.session_state.active_audio_path = destination_path
    st.session_state.show_new_project_uploader = False
    st.session_state.media_uploader_version += 1
    save_material_transcription(project_id, material_id)
    with open(get_project_meta_path(project_id), "w", encoding="utf-8") as project_file:
        json.dump(build_project_meta_from_session(project_id, project_meta), project_file, ensure_ascii=True, indent=2)
    save_project_state()
    return project_id


def activate_quick_split_from_path(source_path, source_name, source_size, max_segment_mb, use_silence=True):
    """Divide un archivo grande ya guardado y activa el primer segmento resultante."""
    if not st.session_state.current_project_id:
        st.warning("Primero crea o abre un proyecto; después divide el audio o video.")
        return []

    project_id = st.session_state.current_project_id
    project_meta = load_project_meta(project_id) or build_project_meta_from_session(project_id)
    if st.session_state.get("active_material_id"):
        save_material_transcription(project_id, st.session_state.active_material_id)

    material_id = build_material_id(
        source_name,
        f"{source_name}-{source_size}-{now_iso()}",
    )
    output_dir = get_project_chunks_dir(project_id, material_id)
    segment_files, metadata = split_large_media_by_size(
        source_path,
        source_name,
        max_segment_mb,
        output_dir,
        use_silence=use_silence,
    )
    if not segment_files:
        raise RuntimeError("No se generaron segmentos a partir del archivo grande.")

    material_record = {
        "id": material_id,
        "name": source_name,
        "path": get_relative_project_path(project_id, source_path),
        "size_bytes": metadata.get("size_bytes", source_size),
        "duration_seconds": metadata.get("duration_seconds"),
        "added": now_iso(),
        "source": "quick_split",
        "split_summary": None,
        "total_segments": len(segment_files),
        "transcribed_segments": 0,
        "status": "pendiente",
        "updated": now_iso(),
    }

    reset_transcription_state()
    st.session_state.active_material_id = material_id
    st.session_state.uploaded_file_id = material_id
    st.session_state.active_audio_name = source_name
    st.session_state.active_audio_path = source_path
    st.session_state.chunks = segment_files
    st.session_state.chunks_prepared = True
    st.session_state.chunk_segment_mins = None
    st.session_state.quick_split_last_summary = {
        "source_name": source_name,
        "segment_count": len(segment_files),
        "max_segment_mb": max_segment_mb,
        "duration": format_duration(metadata.get("duration_seconds")),
        "source_size": format_size_mb(metadata.get("size_bytes", source_size)),
        "output_dir": output_dir,
        "silence_adjustments": metadata.get("silence_adjustments", 0),
        "silence_ranges_count": metadata.get("silence_ranges_count", 0),
        "used_silence_detection": metadata.get("used_silence_detection", False),
    }
    material_record["split_summary"] = st.session_state.quick_split_last_summary.copy()
    project_meta.setdefault("materials", []).append(material_record)
    project_meta["active_material_id"] = material_id
    mark_material_needs_preparation(project_meta)
    st.session_state.show_new_project_uploader = False
    st.session_state.media_uploader_version += 1
    save_material_transcription(project_id, material_id)
    with open(get_project_meta_path(project_id), "w", encoding="utf-8") as project_file:
        json.dump(build_project_meta_from_session(project_id, project_meta), project_file, ensure_ascii=True, indent=2)
    save_project_state()
    return segment_files


def activate_quick_split_segments(uploaded_file, max_segment_mb, use_silence=True):
    """Guarda un archivo grande, lo divide y activa el primer segmento resultante."""
    if not st.session_state.current_project_id:
        st.warning("Primero crea o abre un proyecto; después agrega el audio o video.")
        return []

    source_path = get_project_media_path(
        st.session_state.current_project_id,
        uploaded_file.name,
    )
    save_uploaded_file(uploaded_file, source_path)
    return activate_quick_split_from_path(
        source_path,
        uploaded_file.name,
        uploaded_file.size,
        max_segment_mb,
        use_silence=use_silence,
    )


def save_uploaded_file(uploaded_file, destination_path):
    """Guarda el archivo cargado en una ruta reutilizable."""
    os.makedirs(os.path.dirname(destination_path), exist_ok=True)
    with open(destination_path, "wb") as temp_file:
        temp_file.write(uploaded_file.getbuffer())
    return destination_path


def init_session_state():
    defaults = {
        "uploaded_file_id": None,
        "current_project_id": None,
        "active_material_id": None,
        "active_audio_name": "",
        "active_audio_path": "",
        "chunks": [],
        "chunks_prepared": False,
        "chunk_segment_mins": None,
        "current_chunk_idx": 0,
        "transcript_segments": [],
        "segment_texts": [],
        "current_segment_text": "",
        "memos": [],
        "codes": [],
        "last_transcription": "",
        "autosave_last_saved": "",
        "main_doc_title": "Transcripcion",
        "project_description": "",
        "main_event_date": None,
        "main_transcription_date": datetime.now().strftime("%d/%m/%Y %H:%M"),
        "pending_segment_action": None,
        "show_new_project_uploader": False,
        "show_new_project_form": False,
        "show_project_library_modal": False,
        "show_material_library_modal": False,
        "dialog_project_rename_id": "",
        "dialog_project_delete_id": "",
        "dialog_show_add_material_panel": False,
        "expand_preparation_panel": False,
        "scroll_to_preparation_panel": False,
        "scroll_sidebar_to_top": False,
        "scroll_to_transcription": False,
        "media_uploader_version": 0,
        "language_default_migrated": False,
        "model_default_migrated": False,
        "quick_split_max_mb": WHISPER_FILE_LIMIT_MB,
        "quick_split_last_summary": {},
        "quick_session_mode": False,
        "installer_preview_step": 1,
        "installer_preview_model": DEFAULT_WHISPER_MODEL,
        "installer_context_model": "",
        "installer_install_requested": False,
        "installer_install_error": "",
        "installer_install_messages": [],
        "installer_last_installed_model": "",
    }
    defaults.update(get_default_session_setting_state())
    for key, value in defaults.items():
        if key not in st.session_state:
            st.session_state[key] = value
    if (
        not st.session_state.language_default_migrated
        and st.session_state.sidebar_language_choice == "Detección automática"
    ):
        st.session_state.sidebar_language_choice = "Español"
        st.session_state.language_default_migrated = True
    if (
        not st.session_state.model_default_migrated
        and st.session_state.sidebar_model_choice == "base"
    ):
        st.session_state.sidebar_model_choice = DEFAULT_WHISPER_MODEL
        st.session_state.model_default_migrated = True
    ensure_transcription_stamp()


def reset_transcription_state():
    st.session_state.chunks = []
    st.session_state.chunks_prepared = False
    st.session_state.chunk_segment_mins = None
    st.session_state.current_chunk_idx = 0
    st.session_state.transcript_segments = []
    st.session_state.segment_texts = []
    st.session_state.current_segment_text = ""
    st.session_state.pending_segment_action = None
    st.session_state.memos = []
    st.session_state.codes = []
    st.session_state.last_transcription = ""


def ensure_transcription_stamp():
    """Garantiza una fecha de transcripcion visible y no editable."""
    current_value = st.session_state.get("main_transcription_date", "")
    if not current_value or not str(current_value).strip():
        st.session_state.main_transcription_date = datetime.now().strftime("%d/%m/%Y %H:%M")


def ensure_segment_text_store(total_segments=None):
    """Mantiene un texto independiente por cada segmento navegable."""
    if total_segments is None:
        total_segments = len(st.session_state.get("chunks", []))

    stored = list(st.session_state.get("segment_texts", []))
    legacy_segments = list(st.session_state.get("transcript_segments", []))

    for index, text in enumerate(legacy_segments):
        if index >= len(stored):
            stored.append(text)
        elif not stored[index]:
            stored[index] = text

    current_idx = st.session_state.get("current_chunk_idx", 0)
    current_text = st.session_state.get("current_segment_text", "")
    if current_text and 0 <= current_idx < max(total_segments, len(stored)):
        while len(stored) <= current_idx:
            stored.append("")
        stored[current_idx] = current_text

    if total_segments:
        if len(stored) < total_segments:
            stored.extend([""] * (total_segments - len(stored)))
        elif len(stored) > total_segments:
            stored = stored[:total_segments]

    st.session_state.segment_texts = stored
    st.session_state.transcript_segments = [text for text in stored if text.strip()]
    return stored


def sync_current_segment_text():
    """Guarda el texto visible en el segmento activo antes de navegar."""
    chunks = st.session_state.get("chunks", [])
    if not chunks:
        return
    stored = ensure_segment_text_store(len(chunks))
    current_idx = st.session_state.get("current_chunk_idx", 0)
    if 0 <= current_idx < len(stored):
        stored[current_idx] = st.session_state.get("current_segment_text", "")
    st.session_state.segment_texts = stored
    st.session_state.transcript_segments = [text for text in stored if text.strip()]


def load_segment_text(segment_index):
    """Carga el texto guardado de un segmento al panel de edición."""
    stored = ensure_segment_text_store(len(st.session_state.get("chunks", [])))
    st.session_state.current_chunk_idx = segment_index
    st.session_state.current_segment_text = stored[segment_index] if segment_index < len(stored) else ""


def get_ordered_segment_texts(include_blanks=False):
    """Devuelve los segmentos en orden para exportación y métricas."""
    stored = ensure_segment_text_store(len(st.session_state.get("chunks", [])))
    if include_blanks:
        return stored
    return [text for text in stored if text.strip()]


def render_segment_navigation(chunks):
    """Muestra una linea navegable para saltar entre segmentos."""
    if not chunks:
        return

    stored = ensure_segment_text_store(len(chunks))
    current_idx = st.session_state.get("current_chunk_idx", 0)
    completed_count = sum(1 for text in stored if text.strip())
    st.markdown(
        f"""
        <div class="segment-nav-help">
          Navega por segmentos sin seguir una secuencia lineal. Se guarda el texto actual antes de saltar.
        </div>
        <div class="segment-nav-legend">
          <span>▶ activo</span><span>✓ transcrito</span><span>○ pendiente</span>
          <span>{completed_count}/{len(chunks)} con texto</span>
        </div>
        """,
        unsafe_allow_html=True,
    )

    per_row = 10
    for row_start in range(0, len(chunks), per_row):
        row_items = list(enumerate(chunks[row_start: row_start + per_row], start=row_start))
        cols = st.columns(len(row_items))
        for col, (segment_idx, _) in zip(cols, row_items):
            if segment_idx == current_idx:
                label = f"▶ {segment_idx + 1}"
            elif segment_idx < len(stored) and stored[segment_idx].strip():
                label = f"✓ {segment_idx + 1}"
            else:
                label = f"○ {segment_idx + 1}"
            with col:
                if st.button(label, key=f"segment_nav_{segment_idx}", use_container_width=True):
                    sync_current_segment_text()
                    load_segment_text(segment_idx)
                    save_project_state()
                    st.rerun()








def initialize_project_defaults(uploaded_file=None):
    """Carga valores iniciales cuando no existe un proyecto previo."""
    if uploaded_file is not None:
        st.session_state.main_doc_title = os.path.splitext(uploaded_file.name)[0] or "Entrevista_01"
    st.session_state.main_event_date = None
    st.session_state.main_transcription_date = datetime.now().strftime("%d/%m/%Y %H:%M")
    st.session_state.main_custom_terms = ""
    st.session_state.sidebar_mode = MODE_SEGMENTED
    st.session_state.sidebar_model_choice = DEFAULT_WHISPER_MODEL
    st.session_state.sidebar_language_choice = DEFAULT_TRANSCRIPTION_LANGUAGE
    st.session_state.sidebar_trans_type = "Limpia (Sin muletillas)"
    st.session_state.sidebar_segment_mins = 5
    st.session_state.sidebar_editor_font_size = 16








def list_saved_projects():
    """Lista proyectos locales con metadatos mínimos para el panel."""
    projects = []
    if not os.path.isdir(PROJECTS_DIR):
        return projects

    for project_id in sorted(os.listdir(PROJECTS_DIR)):
        project_dir = get_project_dir(project_id)
        project_meta = load_project_meta(project_id)
        if project_meta and project_meta.get("schema_version") == 2:
            materials = project_meta.get("materials", [])
            active_material = get_active_material_record(project_meta)
            total_segments = sum(int(item.get("total_segments", 0) or 0) for item in materials)
            transcribed_segments = sum(int(item.get("transcribed_segments", 0) or 0) for item in materials)
            projects.append(
                {
                    "id": project_id,
                    "title": project_meta.get("title") or "Proyecto sin título",
                    "description": project_meta.get("description", ""),
                    "date": project_meta.get("event_date", ""),
                    "segments": transcribed_segments,
                    "total_segments": total_segments,
                    "updated": project_meta.get("autosave_last_saved", ""),
                    "has_audio": bool(materials),
                    "active_audio": active_material.get("name", "") if active_material else "",
                    "media_count": len(materials),
                    "progress_ratio": (transcribed_segments / total_segments) if total_segments else 0.0,
                    "mtime": os.path.getmtime(project_dir) if os.path.isdir(project_dir) else 0,
                }
            )
            continue

        state = load_project_state(project_id)
        if not state:
            continue

        projects.append(
            {
                "id": project_id,
                "title": state.get("main_doc_title") or state.get("source_file_name") or "Proyecto sin título",
                "description": state.get("project_description", ""),
                "date": state.get("main_event_date") or "",
                "segments": len(state.get("transcript_segments", [])),
                "total_segments": len(state.get("segment_texts", [])),
                "updated": state.get("autosave_last_saved", ""),
                "has_audio": bool(state.get("source_audio_path") or get_project_audio_source(project_id)),
                "active_audio": state.get("source_file_name", ""),
                "media_count": len(list_project_media(project_id)),
                "progress_ratio": (
                    len(state.get("transcript_segments", [])) / len(state.get("segment_texts", []))
                    if len(state.get("segment_texts", []))
                    else 0.0
                ),
                "mtime": os.path.getmtime(project_dir) if os.path.isdir(project_dir) else 0,
            }
        )

    return sorted(projects, key=lambda item: item["mtime"], reverse=True)


def get_latest_saved_project():
    """Devuelve el proyecto usado más recientemente, si existe."""
    saved_projects = list_saved_projects()
    return saved_projects[0] if saved_projects else None


def get_project_recency_group(project_mtime):
    if not project_mtime:
        return "Anteriores"
    project_dt = datetime.fromtimestamp(project_mtime)
    now = datetime.now()
    if project_dt.date() == now.date():
        return "Hoy"
    if project_dt >= now - timedelta(days=7):
        return "Esta semana"
    return "Anteriores"




def clear_active_project_state():
    """Limpia la mesa actual sin borrar proyectos guardados."""
    st.session_state.current_project_id = None
    st.session_state.active_material_id = None
    st.session_state.uploaded_file_id = None
    st.session_state.active_audio_name = ""
    st.session_state.active_audio_path = ""
    st.session_state.autosave_last_saved = ""
    st.session_state.project_description = ""
    st.session_state.show_new_project_uploader = False
    st.session_state.show_new_project_form = True
    st.session_state.quick_split_last_summary = {}
    st.session_state.quick_session_mode = False
    reset_transcription_state()


def create_project(project_name, project_description):
    """Crea un proyecto local vacío que luego puede recibir audios/videos."""
    if st.session_state.get("current_project_id"):
        save_project_state(write_legacy=True)
    project_id = build_named_project_id(project_name)
    reset_transcription_state()
    st.session_state.current_project_id = project_id
    st.session_state.active_material_id = None
    st.session_state.uploaded_file_id = f"project-{project_id}"
    st.session_state.active_audio_name = ""
    st.session_state.active_audio_path = ""
    st.session_state.main_doc_title = project_name.strip()
    st.session_state.project_description = project_description.strip()
    st.session_state.main_event_date = None
    st.session_state.main_transcription_date = datetime.now().strftime("%d/%m/%Y %H:%M")
    st.session_state.main_custom_terms = ""
    st.session_state.sidebar_mode = MODE_SEGMENTED
    st.session_state.sidebar_model_choice = DEFAULT_WHISPER_MODEL
    st.session_state.sidebar_language_choice = DEFAULT_TRANSCRIPTION_LANGUAGE
    st.session_state.autosave_last_saved = ""
    st.session_state.quick_split_last_summary = {}
    st.session_state.quick_session_mode = False
    st.session_state.show_new_project_form = False
    st.session_state.show_new_project_uploader = True
    save_project_meta(project_id)
    save_project_state()
    return project_id


def create_quick_project():
    """Crea una sesion rapida lista para recibir un material activo."""
    if st.session_state.get("current_project_id"):
        save_project_state(write_legacy=True)
    quick_name = f"Sesion rapida {datetime.now().strftime('%d-%m-%Y %H:%M')}"
    project_id = build_named_project_id(quick_name)
    reset_transcription_state()
    st.session_state.current_project_id = project_id
    st.session_state.active_material_id = None
    st.session_state.uploaded_file_id = f"project-{project_id}"
    st.session_state.active_audio_name = ""
    st.session_state.active_audio_path = ""
    st.session_state.main_doc_title = quick_name
    st.session_state.project_description = ""
    st.session_state.main_event_date = None
    st.session_state.main_transcription_date = datetime.now().strftime("%d/%m/%Y %H:%M")
    st.session_state.main_custom_terms = ""
    st.session_state.sidebar_mode = MODE_SEGMENTED
    st.session_state.sidebar_model_choice = DEFAULT_WHISPER_MODEL
    st.session_state.sidebar_language_choice = DEFAULT_TRANSCRIPTION_LANGUAGE
    st.session_state.autosave_last_saved = ""
    st.session_state.quick_split_last_summary = {}
    st.session_state.quick_session_mode = True
    st.session_state.show_new_project_form = False
    st.session_state.show_new_project_uploader = True
    save_project_meta(project_id)
    save_project_state()
    return project_id


def apply_project_state(project_state):
    """Restaura el estado serializado en session_state."""
    translated_settings = project_settings_to_session_state(project_state.get("settings", {}))
    translated_settings["sidebar_mode"] = normalize_processing_mode(
        project_state.get("sidebar_mode", translated_settings["sidebar_mode"])
    )
    translated_settings["sidebar_model_choice"] = project_state.get(
        "sidebar_model_choice",
        translated_settings["sidebar_model_choice"],
    )
    saved_language = project_state.get(
        "sidebar_language_choice",
        translated_settings["sidebar_language_choice"],
    )
    translated_settings["sidebar_language_choice"] = (
        "Español" if saved_language == "Detección automática" else saved_language
    )
    translated_settings["sidebar_trans_type"] = project_state.get(
        "sidebar_trans_type",
        translated_settings["sidebar_trans_type"],
    )
    translated_settings["sidebar_segment_mins"] = project_state.get(
        "sidebar_segment_mins",
        translated_settings["sidebar_segment_mins"],
    )
    translated_settings["sidebar_editor_font_size"] = project_state.get(
        "sidebar_editor_font_size",
        translated_settings["sidebar_editor_font_size"],
    )
    translated_settings["reading_mode_enabled"] = project_state.get(
        "reading_mode_enabled",
        translated_settings["reading_mode_enabled"],
    )
    for session_key, value in translated_settings.items():
        st.session_state[session_key] = value
    reset_reading_mode_widget_state()
    st.session_state.main_doc_title = project_state.get(
        "main_doc_title",
        st.session_state.main_doc_title,
    )
    st.session_state.quick_session_mode = project_state.get(
        "quick_session_mode",
        False,
    )
    st.session_state.project_description = project_state.get(
        "project_description",
        st.session_state.project_description,
    )
    saved_event_date = project_state.get("main_event_date")
    st.session_state.main_event_date = (
        datetime.fromisoformat(saved_event_date).date()
        if saved_event_date
        else None
    )
    st.session_state.main_transcription_date = project_state.get(
        "main_transcription_date",
        st.session_state.main_transcription_date,
    )
    st.session_state.active_material_id = project_state.get(
        "active_material_id",
        st.session_state.get("active_material_id"),
    )
    st.session_state.main_custom_terms = project_state.get(
        "main_custom_terms",
        st.session_state.main_custom_terms,
    )
    apply_material_transcription_state(project_state)
    st.session_state.quick_split_last_summary = project_state.get("quick_split_last_summary", {})
    st.session_state.autosave_last_saved = project_state.get("autosave_last_saved", "")


def open_existing_project(project_id):
    """Abre un proyecto local previamente guardado."""
    project_state = load_project_state(project_id)
    if not project_state:
        st.warning("No pude abrir este proyecto porque no encontré su ficha guardada.")
        return

    if st.session_state.get("current_project_id") and st.session_state.get("current_project_id") != project_id:
        save_project_state(write_legacy=True)

    reset_transcription_state()
    st.session_state.current_project_id = project_id
    st.session_state.active_material_id = project_state.get("active_material_id")
    st.session_state.uploaded_file_id = project_state.get("uploaded_file_id", st.session_state.active_material_id or project_id)
    st.session_state.active_audio_name = project_state.get("source_file_name", "")
    st.session_state.active_audio_path = project_state.get("source_audio_path", "")
    st.session_state.show_new_project_uploader = not bool(st.session_state.active_audio_path)
    st.session_state.show_new_project_form = False
    apply_project_state(project_state)
    st.rerun()


def save_project_state(write_legacy=False):
    """Guarda automáticamente el avance del proyecto actual."""
    project_id = st.session_state.get("current_project_id")
    active_material_id = st.session_state.get("active_material_id")
    if not project_id:
        return
    if active_material_id:
        save_material_transcription(project_id, active_material_id)
    project_meta = save_project_meta(project_id)
    if write_legacy and project_meta:
        legacy_state = merge_v2_project_state(
            project_meta,
            load_material_transcription(project_id, active_material_id) if active_material_id else empty_material_transcription_state(),
        )
        save_legacy_state_document(project_id, legacy_state)








def restore_project_if_available(uploaded_file):
    """Restaura un proyecto previo o inicializa uno nuevo."""
    project_id = build_project_id(uploaded_file)
    uploaded_file_id = f"{uploaded_file.name}-{uploaded_file.size}"

    if (
        st.session_state.current_project_id == project_id
        and st.session_state.uploaded_file_id == uploaded_file_id
    ):
        return project_id

    reset_transcription_state()
    st.session_state.current_project_id = project_id
    st.session_state.active_material_id = None
    st.session_state.uploaded_file_id = uploaded_file_id
    st.session_state.active_audio_name = uploaded_file.name
    st.session_state.active_audio_path = get_project_audio_path(project_id, uploaded_file.name)
    save_uploaded_file(
        uploaded_file,
        st.session_state.active_audio_path,
    )

    saved_state = load_project_state(project_id)
    if saved_state:
        apply_project_state(saved_state)
    else:
        initialize_project_defaults(uploaded_file)
    st.session_state.show_new_project_uploader = False

    return project_id










def rebuild_active_audio_segments(file_path, segment_mins, project_id, material_id=None):
    """Reconstruye los segmentos del audio activo con la configuración actual."""
    chunks = split_audio(
        file_path,
        segment_mins,
        get_project_chunks_dir(project_id, material_id),
    )
    st.session_state.chunks = chunks
    st.session_state.chunks_prepared = True
    st.session_state.chunk_segment_mins = segment_mins
    st.session_state.current_chunk_idx = 0
    st.session_state.segment_texts = [""] * len(chunks)
    st.session_state.current_segment_text = ""
    save_project_state()
    return chunks


def ensure_project_chunks(audio_path):
    """Restaura o reconstruye los segmentos si el proyecto ya tenía avance."""
    if not st.session_state.chunks_prepared:
        return

    existing_chunks = list_existing_chunks(
        st.session_state.current_project_id,
        st.session_state.get("active_material_id"),
    )
    if existing_chunks:
        st.session_state.chunks = existing_chunks
        ensure_segment_text_store(len(existing_chunks))
        return

    st.session_state.chunks = split_audio(
        audio_path,
        st.session_state.chunk_segment_mins or st.session_state.sidebar_segment_mins,
        get_project_chunks_dir(
            st.session_state.current_project_id,
            st.session_state.get("active_material_id"),
        ),
    )
    ensure_segment_text_store(len(st.session_state.chunks))








def add_markup_runs(paragraph, text):
    """Convierte marcas del editor, incluso combinadas, en formato Word."""
    marker_styles = {
        "**": "bold",
        "__": "underline",
        "==": "highlight",
        "~~": "strike",
    }
    marker_pattern = re.compile(r"(\*\*|__|==|~~)")
    valid_markers = {
        marker
        for marker in marker_styles
        if text.count(marker) >= 2 and text.count(marker) % 2 == 0
    }
    active_styles = {style: False for style in marker_styles.values()}

    for token in marker_pattern.split(text):
        if not token:
            continue
        if token in valid_markers:
            style = marker_styles[token]
            active_styles[style] = not active_styles[style]
            continue

        run = paragraph.add_run(token)
        if active_styles["bold"]:
            run.bold = True
        if active_styles["underline"]:
            run.underline = True
        if active_styles["highlight"]:
            run.font.highlight_color = WD_COLOR_INDEX.YELLOW
        if active_styles["strike"]:
            run.font.strike = True


def add_formatted_transcript(doc, text):
    """Aplica formato simple al cuerpo de la transcripción."""
    for raw_line in text.splitlines():
        line = raw_line.strip()
        if not line:
            continue

        speaker_label, spoken_text = extract_speaker_label(line)
        if speaker_label:
            paragraph = doc.add_paragraph()
            paragraph.add_run(speaker_label).bold = True
            if spoken_text:
                paragraph.add_run(" ")
                add_markup_runs(paragraph, spoken_text)
        else:
            paragraph = doc.add_paragraph()
            add_markup_runs(paragraph, line)


def extract_speaker_label(line):
    """Detecta etiquetas de hablante tipo 'Entrevistadora 1:' al inicio de línea."""
    prefix, separator, remainder = line.partition(":")
    candidate = prefix.strip()
    if not separator or not candidate:
        return None, None

    if len(candidate) > 48 or len(candidate.split()) > 6:
        return None, None

    if not re.fullmatch(r"[A-Za-zÁÉÍÓÚáéíóúÑñ0-9\s.\-_/()]+", candidate):
        return None, None

    return f"{candidate}:", remainder.strip()


def add_word_field(paragraph, instruction):
    """Inserta un campo dinamico de Word, por ejemplo PAGE o NUMPAGES."""
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")

    instr_text = OxmlElement("w:instrText")
    instr_text.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    instr_text.text = instruction

    fld_separate = OxmlElement("w:fldChar")
    fld_separate.set(qn("w:fldCharType"), "separate")

    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")

    run._r.append(fld_begin)
    run._r.append(instr_text)
    run._r.append(fld_separate)
    run._r.append(fld_end)
    run.font.size = Pt(8)
    run.font.italic = True
    run.font.color.rgb = RGBColor(31, 78, 121)
    return run


def enable_docx_field_updates(doc):
    """Pide a Word actualizar campos como PAGE/NUMPAGES al abrir el archivo."""
    settings = doc.settings.element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")


def add_docx_footer(doc):
    """Agrega pie de pagina institucional con pleca azul al documento exportado."""
    footer_text = (
        "Transcripción inmersiva realizada con AudioScript Contextual "
        "Versión Beta 0.9. Desarrollada por Teresa Márquez. "
        "Impulsada por Whisper de OpenIA."
    )
    footer_color = RGBColor(156, 163, 175)

    def format_footer(footer):
        paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        paragraph.alignment = 1
        paragraph_format = paragraph.paragraph_format
        paragraph_format.space_before = Pt(4)
        paragraph_format.space_after = Pt(0)

        p_pr = paragraph._p.get_or_add_pPr()
        p_bdr = p_pr.find(qn("w:pBdr"))
        if p_bdr is None:
            p_bdr = OxmlElement("w:pBdr")
            p_pr.append(p_bdr)
        top = p_bdr.find(qn("w:top"))
        if top is None:
            top = OxmlElement("w:top")
            p_bdr.append(top)
        top.set(qn("w:val"), "single")
        top.set(qn("w:sz"), "12")
        top.set(qn("w:space"), "6")
        top.set(qn("w:color"), "1F4E79")

        paragraph.clear()
        run = paragraph.add_run(footer_text)
        run.font.size = Pt(7.5)
        run.font.italic = False
        run.font.color.rgb = footer_color
        run.add_break()
        page_label = paragraph.add_run("Página ")
        page_label.font.size = Pt(7.5)
        page_label.font.italic = False
        page_label.font.color.rgb = footer_color
        add_word_field(paragraph, " PAGE ")
        page_separator = paragraph.add_run(" de ")
        page_separator.font.size = Pt(7.5)
        page_separator.font.italic = False
        page_separator.font.color.rgb = footer_color
        add_word_field(paragraph, " NUMPAGES ")

    for section in doc.sections:
        section.different_first_page_header_footer = False
        format_footer(section.footer)
        format_footer(section.first_page_footer)
        format_footer(section.even_page_footer)


def build_codes_dataframe():
    """Convierte los códigos guardados en un DataFrame exportable."""
    if not st.session_state.codes:
        return pd.DataFrame()
    return pd.DataFrame(st.session_state.codes)


def render_memo_manager(segment_number):
    """Permite crear, editar y eliminar memos del segmento actual."""
    active_material_id = st.session_state.get("active_material_id", "")
    st.markdown(
        """
        <div class="hover-help-card" tabindex="0" aria-label="Memos de análisis">
          <div class="section-title">Memos de análisis</div>
          <div class="hover-help-card__tip">
            Registra ideas, dudas, conexiones teóricas o decisiones de lectura mientras revisas este material.
          </div>
        </div>
        """,
        unsafe_allow_html=True,
    )

    with st.form(key=f"memo_form_{segment_number}", clear_on_submit=True):
        memo_input = st.text_area(
            "Nueva anotación",
            key=f"memo_input_{segment_number}",
            height=140,
        )
        memo_submitted = st.form_submit_button("Guardar Memo")

    if memo_submitted:
        if memo_input.strip():
            st.session_state.memos.append(
                {
                    "material_id": active_material_id,
                    "segmento": segment_number,
                    "audio": st.session_state.get("active_audio_name", ""),
                    "memo": memo_input.strip(),
                }
            )
            save_project_state()
            st.rerun()
        st.warning("Escribe un memo antes de guardarlo.")

    current_material_memos = [
        (memo_idx, memo)
        for memo_idx, memo in enumerate(st.session_state.memos)
        if memo.get("material_id") == active_material_id
    ]

    if not current_material_memos:
        st.caption("Todavía no hay memos para este material.")
        return

    st.markdown("#### Memos guardados")
    for memo_idx, memo in current_material_memos:
        segment_label = (
            f"Segmento {memo.get('segmento')}"
            if int(memo.get("segmento", 0) or 0) > 0
            else "Documento completo"
        )
        with st.expander(f"Memo {memo_idx + 1} - {segment_label}", expanded=False):
            with st.form(key=f"edit_memo_form_{segment_number}_{memo_idx}"):
                edited_memo = st.text_area(
                    "Editar memo",
                    value=memo["memo"],
                    height=120,
                )
                save_col, delete_col = st.columns(2)
                with save_col:
                    save_memo = st.form_submit_button("Guardar cambios")
                with delete_col:
                    delete_memo = st.form_submit_button("Eliminar memo")

            if save_memo:
                if edited_memo.strip():
                    st.session_state.memos[memo_idx]["memo"] = edited_memo.strip()
                    save_project_state()
                    st.rerun()
                st.warning("El memo no puede quedar vacío.")

            if delete_memo:
                st.session_state.memos.pop(memo_idx)
                save_project_state()
                st.rerun()


def render_code_manager(segment_number):
    """Permite crear, editar y eliminar codificaciones básicas por segmento."""
    active_material_id = st.session_state.get("active_material_id", "")
    st.markdown(
        '<div id="audioscript-code-panel-anchor" style="position:relative; top:-18px;"></div>',
        unsafe_allow_html=True,
    )
    st.markdown(
        """
        <div class="hover-help-card" tabindex="0" aria-label="Pre-codificación o codificación gruesa">
          <div class="section-title">Pre-codificación o codificación gruesa</div>
          <div class="hover-help-card__tip">
            Pega una frase seleccionada, nómbrala con un código provisional y agrega una nota analítica si hace falta.
          </div>
        </div>
        """,
        unsafe_allow_html=True,
    )

    with st.form(key=f"code_form_{segment_number}", clear_on_submit=True):
        quote_text = st.text_area(
            "Cita o frase a codificar",
            key=f"code_quote_{segment_number}",
            height=120,
            placeholder="Pega aquí la frase exacta del segmento transcrito.",
        )
        code_label = st.text_input(
            "Código",
            key=f"code_label_{segment_number}",
            placeholder="Ej: identidad, conflicto, precariedad...",
        )
        code_note = st.text_area(
            "Nota analítica opcional",
            key=f"code_note_{segment_number}",
            height=100,
        )
        code_submitted = st.form_submit_button("Guardar código")

    if code_submitted:
        if quote_text.strip() and code_label.strip():
            st.session_state.codes.append(
                {
                    "material_id": active_material_id,
                    "segmento": segment_number,
                    "audio": st.session_state.get("active_audio_name", ""),
                    "cita": quote_text.strip(),
                    "codigo": code_label.strip(),
                    "nota": code_note.strip(),
                    "fecha_registro": datetime.now().strftime("%d/%m/%Y %H:%M:%S"),
                }
            )
            save_project_state()
            st.rerun()
        st.warning("Completa al menos la cita y el nombre del código.")

    current_material_codes = [
        (code_idx, code_item)
        for code_idx, code_item in enumerate(st.session_state.codes)
        if code_item.get("material_id") == active_material_id
    ]

    if not current_material_codes:
        st.caption("Todavía no hay códigos para este material.")
        return

    st.markdown("#### Códigos guardados")
    for code_idx, code_item in current_material_codes:
        segment_label = (
            f"Segmento {code_item.get('segmento')}"
            if int(code_item.get("segmento", 0) or 0) > 0
            else "Documento completo"
        )
        with st.expander(f"Código {code_idx + 1} - {segment_label}: {code_item['codigo']}", expanded=False):
            with st.form(key=f"edit_code_form_{segment_number}_{code_idx}"):
                edited_quote = st.text_area(
                    "Cita",
                    value=code_item["cita"],
                    height=120,
                )
                edited_label = st.text_input(
                    "Código",
                    value=code_item["codigo"],
                )
                edited_note = st.text_area(
                    "Nota analítica",
                    value=code_item.get("nota", ""),
                    height=100,
                )
                save_code_col, delete_code_col = st.columns(2)
                with save_code_col:
                    save_code = st.form_submit_button("Guardar cambios")
                with delete_code_col:
                    delete_code = st.form_submit_button("Eliminar código")

            if save_code:
                if edited_quote.strip() and edited_label.strip():
                    st.session_state.codes[code_idx] = {
                        **st.session_state.codes[code_idx],
                        "cita": edited_quote.strip(),
                        "codigo": edited_label.strip(),
                        "nota": edited_note.strip(),
                    }
                    save_project_state()
                    st.rerun()
                st.warning("La cita y el código son obligatorios.")

            if delete_code:
                st.session_state.codes.pop(code_idx)
                save_project_state()
                st.rerun()






def save_as_docx(text, title, event_date, transcription_date, project_context, memos_df, codes_df):
    """Genera un archivo .docx con metadatos, cuerpo, memos y códigos."""
    doc = Document()
    enable_docx_field_updates(doc)
    doc.add_heading(title or "Transcripción", 0)
    safe_event_date = event_date or "No especificada"
    safe_transcription_date = transcription_date or datetime.now().strftime("%d/%m/%Y %H:%M")
    safe_project_context = (project_context or "").strip()

    meta_paragraph = doc.add_paragraph()
    meta_paragraph.add_run("Fecha del evento: ").bold = True
    meta_paragraph.add_run(f"{safe_event_date}\n")
    meta_paragraph.add_run("Fecha de transcripción: ").bold = True
    meta_paragraph.add_run(f"{safe_transcription_date}\n")

    context_heading = doc.add_paragraph()
    context_heading.add_run("Objetivo del evento y contexto").bold = True
    context_paragraph = doc.add_paragraph()
    context_paragraph.add_run(
        safe_project_context or "No especificado."
    )

    doc.add_paragraph("-" * 30)
    add_formatted_transcript(doc, text)

    if not memos_df.empty:
        doc.add_page_break()
        doc.add_heading("Notas de Investigación (Memos)", level=1)
        for _, row in memos_df.iterrows():
            paragraph = doc.add_paragraph()
            paragraph.add_run(f"Segmento {row['segmento']}: ").bold = True
            paragraph.add_run(str(row["memo"]))

    if not codes_df.empty:
        doc.add_page_break()
        doc.add_heading("Codificación", level=1)
        for _, row in codes_df.iterrows():
            paragraph = doc.add_paragraph()
            paragraph.add_run(f"Segmento {row['segmento']} | Código: ").bold = True
            paragraph.add_run(f"{row['codigo']}\n")
            paragraph.add_run("Cita: ").bold = True
            paragraph.add_run(f"{row['cita']}\n")
            if str(row.get("nota", "")).strip():
                paragraph.add_run("Nota: ").bold = True
                paragraph.add_run(f"{row['nota']}\n")

    try:
        add_docx_footer(doc)
    except Exception:
        # Some frozen python-docx builds cannot resolve their footer template.
        # The transcript is still valuable, so export it without the footer.
        LOGGER.exception("No se pudo agregar el pie de pagina al DOCX")

    temp_docx = tempfile.NamedTemporaryFile(
        delete=False,
        suffix=".docx",
        dir=TEMP_DIR,
    )
    temp_docx_path = temp_docx.name
    temp_docx.close()
    doc.save(temp_docx_path)
    return temp_docx_path


def dataframe_to_xlsx_bytes(dataframes_by_sheet):
    """Convierte uno o varios DataFrames a un archivo XLSX en memoria."""
    excel_engine = None
    if importlib.util.find_spec("openpyxl"):
        excel_engine = "openpyxl"
    elif importlib.util.find_spec("xlsxwriter"):
        excel_engine = "xlsxwriter"
    if excel_engine is None:
        raise RuntimeError("No hay motor Excel disponible para exportar a XLSX.")

    output = io.BytesIO()
    with pd.ExcelWriter(output, engine=excel_engine) as writer:
        for sheet_name, dataframe in dataframes_by_sheet.items():
            dataframe.to_excel(writer, index=False, sheet_name=sheet_name[:31] or "Datos")
    output.seek(0)
    return output.getvalue()


def build_memos_export_dataframe(memos_df):
    """Normaliza memos para exportación en hoja de cálculo."""
    if memos_df.empty:
        return pd.DataFrame(columns=["Memo", "Segmento", "Audio"])
    export_df = memos_df.copy()
    export_df = export_df.rename(
        columns={
            "memo": "Memo",
            "segmento": "Segmento",
            "audio": "Audio",
        }
    )
    return export_df[["Memo", "Segmento", "Audio"]]


def build_codes_import_dataframe(codes_df):
    """Crea una hoja mínima compatible con importadores que esperan Código."""
    if codes_df.empty:
        return pd.DataFrame(columns=["Código"])

    rows = []
    seen = set()
    for _, row in codes_df.iterrows():
        code_label = str(row.get("codigo", "")).strip()
        if not code_label:
            continue
        fingerprint = code_label.lower()
        if fingerprint in seen:
            continue
        seen.add(fingerprint)
        rows.append({"Código": code_label})
    return pd.DataFrame(rows, columns=["Código"])


def build_codes_detail_dataframe(codes_df):
    """Conserva detalle completo de codificación en una segunda hoja."""
    if codes_df.empty:
        return pd.DataFrame(
            columns=["Código", "Memo", "Cita", "Segmento", "Audio", "Fecha de registro"]
        )

    export_df = codes_df.copy()
    export_df["Memo"] = export_df.get("nota", "").fillna("")
    export_df = export_df.rename(
        columns={
            "codigo": "Código",
            "cita": "Cita",
            "segmento": "Segmento",
            "audio": "Audio",
            "fecha_registro": "Fecha de registro",
        }
    )
    return export_df[["Código", "Memo", "Cita", "Segmento", "Audio", "Fecha de registro"]]


def build_refi_qda_codebook(title, codes_df):
    """Genera un archivo QDC basico compatible con REFI-QDA Codebook."""
    ET.register_namespace("", "urn:QDA-XML:codebook:1.0")
    root = ET.Element(
        "CodeBook",
        attrib={
            "xmlns": "urn:QDA-XML:codebook:1.0",
            "origin": "AudioScript Contextual",
            "creatingUser": "AudioScript",
            "creationDateTime": datetime.now().isoformat(timespec="seconds"),
            "guid": str(uuid.uuid4()),
        },
    )

    codes_el = ET.SubElement(root, "Codes")
    codes_detail_df = build_codes_detail_dataframe(codes_df)
    seen = set()
    for _, row in codes_detail_df.iterrows():
        code_label = str(row.get("Código", "")).strip()
        memo_text = str(row.get("Memo", "")).strip() or str(row.get("Cita", "")).strip()
        if not code_label or code_label.lower() in seen:
            continue
        seen.add(code_label.lower())
        code_el = ET.SubElement(
            codes_el,
            "Code",
            attrib={
                "name": code_label,
                "guid": str(uuid.uuid4()),
                "isCodable": "true",
            },
        )
        description_el = ET.SubElement(code_el, "Description")
        description_el.text = memo_text or f"Exportado desde {title or 'AudioScript Contextual'}."

    xml_bytes = ET.tostring(root, encoding="utf-8", xml_declaration=True)
    return xml_bytes


def render_downloads(text, title, event_date, transcription_date, memos_df=None, codes_df=None):
    """Renderiza botones de descarga para texto, docx, memos y códigos."""
    if memos_df is None:
        memos_df = pd.DataFrame(st.session_state.memos)
    if codes_df is None:
        codes_df = build_codes_dataframe()
    excel_export_available = bool(
        importlib.util.find_spec("openpyxl") or importlib.util.find_spec("xlsxwriter")
    )

    col1, col2, col3, col4, col5 = st.columns(5)
    with col1:
        st.download_button(
            "Descargar .TXT",
            data=text,
            file_name=f"{title}.txt",
            mime="text/plain",
        )
    with col2:
        try:
            doc_path = save_as_docx(
                text,
                title,
                event_date,
                transcription_date,
                st.session_state.get("project_description", ""),
                memos_df,
                codes_df,
            )
            with open(doc_path, "rb") as doc_file:
                st.download_button(
                    "Descargar .DOCX",
                    data=doc_file.read(),
                    file_name=f"{title}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
        except Exception:
            LOGGER.exception("No se pudo generar la exportacion DOCX")
            st.button("DOCX temporalmente no disponible", disabled=True)
            st.caption(
                "La transcripción sigue disponible. El detalle quedó registrado "
                "en el archivo de diagnóstico de AudioScript."
            )
    with col3:
        if memos_df.empty:
            st.button("No hay memos para exportar", disabled=True)
        elif not excel_export_available:
            st.button("Memos (XLSX no disponible)", disabled=True)
        else:
            xlsx_data = dataframe_to_xlsx_bytes(
                {"Memos": build_memos_export_dataframe(memos_df)}
            )
            st.download_button(
                "Descargar Memos (XLSX)",
                data=xlsx_data,
                file_name=f"{title}_memos.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            )
    with col4:
        if codes_df.empty:
            st.button("No hay códigos para exportar", disabled=True)
        elif not excel_export_available:
            st.button("Códigos (XLSX no disponible)", disabled=True)
        else:
            xlsx_data = dataframe_to_xlsx_bytes(
                {
                    "Diccionario": build_codes_import_dataframe(codes_df),
                    "Codificaciones": build_codes_detail_dataframe(codes_df),
                }
            )
            st.download_button(
                "Descargar Códigos (XLSX)",
                data=xlsx_data,
                file_name=f"{title}_codigos.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            )
    with col5:
        if codes_df.empty:
            st.button("No hay codebook QDC", disabled=True)
        else:
            qdc_data = build_refi_qda_codebook(title, codes_df)
            st.download_button(
                "Exportar Codebook (QDC)",
                data=qdc_data,
                file_name=f"{title}_codebook.qdc",
                mime="application/xml",
            )
    if not excel_export_available and (not memos_df.empty or not codes_df.empty):
        st.caption("La exportación XLSX requiere instalar `openpyxl` o `xlsxwriter` en esta versión de la app.")


def render_quick_split_controls(source_label, source_size, key_prefix):
    """Controles básicos de Fase 1 para dividir archivos grandes."""
    st.markdown("#### Quick Split")
    st.caption(
        f"{source_label} pesa {format_size_mb(source_size)}. AudioScript puede crear segmentos temporales para transcribirlos por partes."
    )
    size_choice = st.radio(
        "Tamaño máximo por segmento",
        [25, 50, 100, 200, 500, "Personalizado"],
        index=3,
        horizontal=True,
        key=f"{key_prefix}_quick_split_size_choice",
    )
    if size_choice == "Personalizado":
        max_mb = st.number_input(
            "Tamaño personalizado en MB",
            min_value=10,
            max_value=900,
            value=int(st.session_state.quick_split_max_mb),
            step=10,
            key=f"{key_prefix}_quick_split_custom_size",
        )
    else:
        max_mb = int(size_choice)
    st.session_state.quick_split_max_mb = int(max_mb)
    use_silence = st.checkbox(
        "Buscar silencio antes del corte",
        value=True,
        key=f"{key_prefix}_quick_split_use_silence",
        help="Evita cortar palabras a la mitad detectando pausas naturales cerca del punto de corte.",
    )
    estimated_segments = max(1, int((source_size + (int(max_mb) * BYTES_PER_MB) - 1) // (int(max_mb) * BYTES_PER_MB)))
    silence_note = (
        "Se buscarán pausas naturales cercanas a cada corte."
        if use_silence
        else "Los cortes se harán por tamaño/duración estimada, sin buscar pausas."
    )
    st.info(f"Se crearán aproximadamente {estimated_segments} segmento(s). {silence_note}")
    render_quick_split_preview(source_label, source_size, int(max_mb), use_silence, key_prefix)
    return int(max_mb), use_silence


def trigger_sidebar_flow_step(step_key):
    if step_key == "project":
        st.session_state.show_project_library_modal = True
    elif step_key == "material":
        if st.session_state.get("current_project_id"):
            st.session_state.show_material_library_modal = True
    elif step_key == "preparation":
        if st.session_state.get("current_project_id") and st.session_state.get("active_audio_path"):
            st.session_state.expand_preparation_panel = True
            st.session_state.scroll_to_preparation_panel = True
    elif step_key == "transcribe":
        if st.session_state.get("active_audio_path"):
            st.session_state.scroll_to_transcription = True
    st.rerun()


def render_sidebar_flow_steps(project_meta=None):
    """Muestra la ruta de preparación y permite saltar a cada paso."""
    has_project = bool(st.session_state.get("current_project_id"))
    has_material = bool(st.session_state.get("active_audio_path"))
    preparation_done = has_project and has_material and is_project_preparation_confirmed(project_meta)
    transcription_started = has_project and has_material and has_transcription_started(project_meta)
    steps = [
        ("project", "Proyecto", "Abre o crea el contexto de trabajo.", has_project, not has_project),
        ("material", "Material", "Carga el audio o video.", has_material, has_project and not has_material),
        ("preparation", "Preparación", "Confirma modo, idioma y modelo.", preparation_done, has_project and has_material and not preparation_done),
        ("transcribe", "Transcribir", "Inicia cuando todo esté listo.", transcription_started, preparation_done and not transcription_started),
    ]
    html_steps = []
    for index, (_, title, detail, done, active) in enumerate(steps, start=1):
        class_name = "sidebar-flow-step"
        if done:
            class_name += " sidebar-flow-step--done"
        elif active:
            class_name += " sidebar-flow-step--active"
        marker = "✓" if done else str(index)
        html_steps.append(
            f'<div class="{class_name}">'
            f'<div class="sidebar-flow-step__num">{marker}</div>'
            f'<div><b>{html.escape(title)}</b><span>{html.escape(detail)}</span></div>'
            '</div>'
        )
    st.markdown(
        '<div class="sidebar-flow-steps">' + "".join(html_steps) + '</div>',
        unsafe_allow_html=True,
    )


def is_project_preparation_confirmed(project_meta=None):
    project_id = st.session_state.get("current_project_id")
    if not project_id:
        return False

    project_meta = project_meta or load_project_meta(project_id)
    if not project_meta:
        return False

    stored_settings = project_meta.get("settings", {})
    current_settings = build_project_meta_from_session(project_id, project_meta).get("settings", {})
    tracked_keys = [
        "mode",
        "model",
        "language",
        "trans_type",
        "segment_mins",
        "editor_font_size",
        "custom_terms",
        "reading_mode",
    ]
    settings_match = all(
        stored_settings.get(key) == current_settings.get(key)
        for key in tracked_keys
    )
    return bool(stored_settings.get("prep_confirmed")) and settings_match


def has_transcription_started(project_meta=None):
    project_id = st.session_state.get("current_project_id")
    if not project_id:
        return False

    project_meta = project_meta or load_project_meta(project_id)
    if not project_meta:
        return False

    settings = project_meta.get("settings", {})
    if settings.get("transcription_started"):
        return True

    active_material = get_active_material_record(project_meta)
    if active_material:
        if int(active_material.get("transcribed_segments", 0) or 0) > 0:
            return True
        material_state = load_material_transcription(project_id, active_material.get("id", ""))
        if material_state.get("last_transcription", "").strip():
            return True
        if any(str(text).strip() for text in material_state.get("segment_texts", [])):
            return True
    return False


def confirm_project_preparation():
    project_id = st.session_state.get("current_project_id")
    if not project_id:
        return
    project_meta = load_project_meta(project_id) or build_project_meta_from_session(project_id)
    settings = project_meta.setdefault("settings", get_default_project_settings())
    settings.update(build_project_meta_from_session(project_id, project_meta).get("settings", {}))
    settings["prep_confirmed"] = True
    with open(get_project_meta_path(project_id), "w", encoding="utf-8") as project_file:
        json.dump(build_project_meta_from_session(project_id, project_meta), project_file, ensure_ascii=True, indent=2)


def mark_transcription_started():
    project_id = st.session_state.get("current_project_id")
    if not project_id:
        return
    project_meta = load_project_meta(project_id) or build_project_meta_from_session(project_id)
    settings = project_meta.setdefault("settings", get_default_project_settings())
    settings.update(build_project_meta_from_session(project_id, project_meta).get("settings", {}))
    settings["transcription_started"] = True
    with open(get_project_meta_path(project_id), "w", encoding="utf-8") as project_file:
        json.dump(build_project_meta_from_session(project_id, project_meta), project_file, ensure_ascii=True, indent=2)


def get_contextual_sidebar_action(project_meta=None):
    project_id = st.session_state.get("current_project_id")
    if not project_id:
        return {
            "key": "project_library",
            "label": "+ Crear o abrir proyecto",
            "detail": "Abre un proyecto o crea uno para empezar a trabajar.",
        }

    project_meta = project_meta or load_project_meta(project_id)
    materials = project_meta.get("materials", []) if project_meta else []
    active_material_id = st.session_state.get("active_material_id")
    resolved_active_material = (
        get_material_record(project_meta, active_material_id)
        if project_meta and active_material_id
        else get_active_material_record(project_meta)
    )
    has_active_material = bool(resolved_active_material and st.session_state.get("active_audio_path"))
    if materials and not has_active_material:
        return {
            "key": "materials",
            "label": f"Elegir material ({len(materials)})",
            "detail": "Abre la biblioteca de materiales del proyecto para activarlo o cargar uno nuevo.",
        }
    if not materials:
        return {
            "key": "materials",
            "label": "Cargar audio o video",
            "detail": "Vincula el primer material del proyecto desde un único panel.",
        }
    if not is_project_preparation_confirmed(project_meta):
        return {
            "key": "preparation",
            "label": "Revisar preparación",
            "detail": "Confirma modo, idioma, segmentos y modelo antes de transcribir.",
        }
    return {
        "key": "transcribe",
        "label": "Transcribir material activo",
        "detail": "La mesa central ya está lista para iniciar o continuar la transcripción.",
    }


def render_sidebar_primary_action(action):
    class_name = "sidebar-primary-action"
    if action["key"] == "transcribe":
        class_name += " sidebar-primary-action--hot"
    st.markdown(
        f"""
        <div class="{class_name}">
          <strong>{html.escape(action["label"])}</strong>
          <span>{html.escape(action["detail"])}</span>
        </div>
        """,
        unsafe_allow_html=True,
    )
    if st.button(action["label"], key=f"sidebar_contextual_action_{action['key']}", type="primary", use_container_width=True):
        if action["key"] == "project_library":
            st.session_state.show_project_library_modal = True
        elif action["key"] == "materials":
            st.session_state.show_material_library_modal = True
        elif action["key"] == "preparation":
            st.session_state.expand_preparation_panel = True
            st.session_state.scroll_to_preparation_panel = True
        elif action["key"] == "transcribe":
            st.session_state.scroll_to_transcription = True
        st.rerun()


def render_preparation_summary(project_meta=None):
    current_project_id = st.session_state.get("current_project_id")
    if not current_project_id:
        preparation_status = "Pendiente"
        transcription_status = "Aún no inicia"
    else:
        project_meta = project_meta or load_project_meta(current_project_id)
        preparation_status = "Confirmada" if is_project_preparation_confirmed(project_meta) else "Pendiente"
        transcription_status = "Iniciada" if has_transcription_started(project_meta) else "Aún no inicia"
    mode_short = normalize_processing_mode(st.session_state.sidebar_mode)
    st.markdown(
        '<div class="sidebar-compact-settings">'
        f'<div class="sidebar-setting-line"><b>Modo</b><span>{html.escape(mode_short)}</span></div>'
        f'<div class="sidebar-setting-line"><b>Idioma</b><span>{html.escape(st.session_state.sidebar_language_choice)}</span></div>'
        f'<div class="sidebar-setting-line"><b>Segmentos</b><span>{int(st.session_state.sidebar_segment_mins)} min</span></div>'
        f'<div class="sidebar-setting-line"><b>Modelo</b><span>{html.escape(st.session_state.sidebar_model_choice)}</span></div>'
        f'<div class="sidebar-setting-line"><b>Formato</b><span>{html.escape("Limpia" if "Limpia" in st.session_state.sidebar_trans_type else "Verbatim")}</span></div>'
        '</div>',
        unsafe_allow_html=True,
    )




def reset_reading_mode_widget_state():
    widget_keys = [
        "complete_reading_mode_toggle",
        "final_reading_mode_toggle",
    ]
    widget_keys.extend(
        key
        for key in list(st.session_state.keys())
        if key.startswith("segment_reading_mode_toggle_")
    )
    for key in widget_keys:
        st.session_state.pop(key, None)


def sync_reading_mode_toggle(widget_key):
    st.session_state.reading_mode_enabled = bool(st.session_state.get(widget_key, False))
    save_project_state()


def render_reading_mode_toggle(widget_key, help_text=""):
    current_value = bool(st.session_state.get("reading_mode_enabled", False))
    if widget_key not in st.session_state:
        st.session_state[widget_key] = current_value
    effective_help = help_text or (
        "Siguiendo recomendaciones de estudios de tipografia y seguimiento ocular."
    )
    st.toggle(
        "Modo lectura",
        key=widget_key,
        help=effective_help,
        on_change=sync_reading_mode_toggle,
        args=(widget_key,),
    )
    return bool(st.session_state.get(widget_key, current_value))


def render_reading_mode_hint():
    return


def resolve_installer_target_model(required_model=None):
    normalized_required = str(required_model or DEFAULT_WHISPER_MODEL).strip().lower()
    if normalized_required in GUIDED_INSTALL_MODEL_OPTIONS:
        selected_model = str(
            st.session_state.get("installer_preview_model", normalized_required)
        ).strip().lower()
        if selected_model in GUIDED_INSTALL_MODEL_OPTIONS:
            return selected_model
        return normalized_required
    return normalized_required or DEFAULT_WHISPER_MODEL


def installer_onboarding_completed():
    return os.path.isfile(INSTALLER_COMPLETED_MARKER)


def mark_installer_onboarding_completed():
    os.makedirs(os.path.dirname(INSTALLER_COMPLETED_MARKER), exist_ok=True)
    with open(INSTALLER_COMPLETED_MARKER, "w", encoding="utf-8") as marker_file:
        marker_file.write(datetime.now().isoformat())


def installer_flow_is_active(required_model=None):
    if not installer_onboarding_completed():
        return True
    return bool(
        st.session_state.get("installer_preview_step") == 5
        and st.session_state.get("installer_last_installed_model")
    )


def render_installer_gate_page(message, required_model=None):
    """Presenta una primera instalación limpia, sin exponer el resto de la app."""
    st.markdown(
        """
        <style>
        section[data-testid="stSidebar"],
        div[data-testid="stSidebarCollapsedControl"],
        div[data-testid="collapsedControl"],
        div[data-testid="stSidebarCollapseButton"],
        button[aria-label*="sidebar" i],
        #audioscript-sidebar-rescue {
            display: none !important;
            visibility: hidden !important;
            opacity: 0 !important;
            pointer-events: none !important;
        }
        .block-container {
            padding-top: 1.2rem !important;
        }
        </style>
        """,
        unsafe_allow_html=True,
    )
    render_app_header()
    st.markdown(
        f"""
        <section class="installer-gate-shell">
          <div class="installer-gate-shell__intro">
            <div class="installer-gate-shell__eyebrow">Primera apertura en esta Mac</div>
            <h2>Preparación inicial antes de transcribir</h2>
            <p>{html.escape(message)}</p>
          </div>
        </section>
        """,
        unsafe_allow_html=True,
    )
    render_whisper_installer_flow(required_model)


def render_whisper_installer_flow(required_model=None):
    """Flujo real de instalación inicial del modelo Whisper para la beta Mac."""
    required_model = str(required_model or DEFAULT_WHISPER_MODEL).strip().lower()
    if st.session_state.get("installer_context_model") != required_model:
        st.session_state.installer_context_model = required_model
        st.session_state.installer_install_requested = False
        st.session_state.installer_install_error = ""
        st.session_state.installer_install_messages = []
        st.session_state.installer_preview_step = 1
        if required_model in GUIDED_INSTALL_MODEL_OPTIONS:
            st.session_state.installer_preview_model = required_model

    installer_target = resolve_installer_target_model(required_model)
    install_info = get_whisper_model_installation_info(installer_target)

    st.markdown(
        f"""
        <div class="installer-preview-shell">
          <div class="project-context-card installer-preview-context">
            <div class="project-context-card__title">Instalación inicial del motor local</div>
            <div class="project-context-card__hint">
              AudioScript instalará una sola vez el modelo <b>{html.escape(installer_target.title())}</b>
              dentro de tu carpeta local. Después, las transcripciones se ejecutarán sin enviar audios a servidores.
            </div>
            <div class="installer-preview-context__meta">Beta Mac Apple Silicon · Instalación única · Funcionamiento local posterior</div>
          </div>
        </div>
        """,
        unsafe_allow_html=True,
    )

    step_labels = [
        "1. Bienvenida",
        "2. Modelos",
        "3. Antes de instalar",
        "4. Progreso",
        "5. Listo",
    ]
    tracker_markup = "".join(
        (
            f'<div class="installer-progress-step {"is-current" if st.session_state.installer_preview_step == step_index else ""} '
            f'{"is-complete" if st.session_state.installer_preview_step > step_index else ""}">'
            f"<span>{html.escape(label)}</span>"
            "</div>"
        )
        for step_index, label in enumerate(step_labels, start=1)
    )
    st.markdown(
        f'<div class="installer-progress-tracker">{tracker_markup}</div>',
        unsafe_allow_html=True,
    )

    if st.session_state.installer_preview_step >= 2:
        st.markdown(
            '<div class="installer-preview-intro">Elige el modelo con el que quedará preparada esta beta en la Mac de la tester.</div>',
            unsafe_allow_html=True,
        )
        st.markdown('<div class="installer-preview-model-switch">', unsafe_allow_html=True)
        model_cols = st.columns(2, gap="small")
        for model_name, col in (("medium", model_cols[0]), ("large", model_cols[1])):
            button_type = "primary" if st.session_state.installer_preview_model == model_name else "secondary"
            if col.button(
                f"{'Usar' if st.session_state.installer_preview_model != model_name else 'Viendo'} {model_name.title()}",
                key=f"installer_preview_model_{model_name}",
                type=button_type,
                use_container_width=True,
            ):
                st.session_state.installer_preview_model = model_name
                st.rerun()
        st.markdown("</div>", unsafe_allow_html=True)

    render_beta_installer_preview(
        st.session_state.installer_preview_step,
        installer_target,
    )

    if st.session_state.installer_preview_step == 4:
        progress_placeholder = st.empty()
        status_placeholder = st.empty()

        def update_install_status(stage, message):
            progress_map = {
                "prepare": 18,
                "download": 55,
                "verify": 82,
                "complete": 100,
            }
            messages = list(st.session_state.get("installer_install_messages", []))
            if not messages or messages[-1] != message:
                messages.append(message)
                st.session_state.installer_install_messages = messages
            progress_placeholder.progress(progress_map.get(stage, 10), text=message)
            status_placeholder.markdown(
                "\n".join(f"- {html.escape(item)}" for item in messages)
            )

        if st.session_state.get("installer_install_requested"):
            try:
                ensure_whisper_model_installed(installer_target, update_install_status)
                load_whisper_model.clear()
                st.session_state.sidebar_model_choice = installer_target
                st.session_state.installer_install_requested = False
                st.session_state.installer_install_error = ""
                st.session_state.installer_last_installed_model = installer_target
                st.session_state.installer_preview_step = 5
                if st.session_state.get("current_project_id"):
                    save_project_state()
                st.rerun()
            except Exception as exc:
                st.session_state.installer_install_requested = False
                st.session_state.installer_install_error = str(exc)
                progress_placeholder.empty()
                status_placeholder.empty()

        if st.session_state.get("installer_install_error"):
            st.error(f"No fue posible instalar el modelo: {st.session_state.installer_install_error}")
        elif st.session_state.get("installer_install_messages"):
            st.success("Instalación en curso o lista para finalizar.")
            st.markdown(
                "\n".join(
                    f"- {html.escape(item)}"
                    for item in st.session_state.get("installer_install_messages", [])
                )
            )

    if st.session_state.installer_preview_step == 5:
        installed_label = st.session_state.get("installer_last_installed_model") or installer_target
        installed_path = install_info.get("path") or "la carpeta local de AudioScript"
        st.success(
            f"El modelo {installed_label.title()} quedó instalado en {installed_path}."
        )

    st.markdown('<div class="installer-preview-transport">', unsafe_allow_html=True)
    transport_cols = st.columns([1, 1], gap="small")
    current_step = st.session_state.installer_preview_step
    back_disabled = current_step <= 1
    if transport_cols[0].button(
        "Atrás",
        key="installer_preview_back",
        disabled=back_disabled,
        use_container_width=True,
    ):
        st.session_state.installer_preview_step = max(1, current_step - 1)
        st.rerun()

    forward_label = {
        1: "Continuar",
        2: "Continuar",
        3: "Comenzar instalación",
        4: "Reintentar instalación",
        5: "Entrar a AudioScript",
    }.get(current_step, "Continuar")
    if transport_cols[1].button(
        forward_label,
        key="installer_preview_forward",
        type="primary",
        use_container_width=True,
    ):
        if current_step == 3:
            st.session_state.installer_install_requested = True
            st.session_state.installer_install_error = ""
            st.session_state.installer_install_messages = []
            st.session_state.installer_preview_step = 4
        elif current_step == 4:
            st.session_state.installer_install_requested = True
            st.session_state.installer_install_error = ""
            st.session_state.installer_install_messages = []
        elif current_step >= 5:
            mark_installer_onboarding_completed()
            st.session_state.installer_preview_step = 1
            st.session_state.installer_last_installed_model = ""
            st.session_state.installer_context_model = ""
        else:
            st.session_state.installer_preview_step = current_step + 1
        st.rerun()
    st.markdown("</div>", unsafe_allow_html=True)
    st.caption(
        "La descarga ocurre una sola vez. Después, el modelo queda guardado en la carpeta local de AudioScript y la transcripción corre sin conexión."
    )


@st.dialog("Biblioteca", width="small")
def render_project_library_dialog():
    st.session_state.show_project_library_modal = True
    saved_projects = list_saved_projects()

    st.markdown('<div class="dialog-title">Biblioteca de proyectos</div>', unsafe_allow_html=True)

    project_search = st.text_input(
        "Buscar por nombre del proyecto",
        value="",
        placeholder="Buscar por nombre del proyecto...",
        key="dialog_project_search",
        label_visibility="collapsed",
    ).strip().lower()

    create_col, create_btn_col = st.columns([5, 1])
    with create_col:
        project_name = st.text_input(
            "Nombre de un proyecto nuevo",
            placeholder="Nombre de un proyecto nuevo...",
            key="dialog_new_project_name",
            label_visibility="collapsed",
        )
    with create_btn_col:
        create_submitted = st.button("Crear", key="dialog_create_project_inline", type="primary", use_container_width=True)
    if create_submitted:
        if project_name.strip():
            st.session_state.show_project_library_modal = False
            create_project(project_name, "")
            st.rerun()
        st.warning("Escribe un nombre para crear el proyecto.")

    if not saved_projects:
        st.caption("Todavía no hay proyectos guardados.")
        return

    filtered_projects = []
    for project in saved_projects:
        haystack = " ".join(
            [
                str(project.get("title", "")),
                str(project.get("description", "")),
                str(project.get("active_audio", "")),
            ]
        ).lower()
        if not project_search or project_search in haystack:
            filtered_projects.append(project)

    if not filtered_projects:
        st.caption("No encontré proyectos con ese criterio.")
        return

    grouped_projects = {"Hoy": [], "Esta semana": [], "Anteriores": []}
    for project in filtered_projects:
        grouped_projects[get_project_recency_group(project.get("mtime"))].append(project)

    for group_label in ["Hoy", "Esta semana", "Anteriores"]:
        if not grouped_projects[group_label]:
            continue
        st.markdown(f'<div class="dialog-section-label">{group_label}</div>', unsafe_allow_html=True)
        for project in grouped_projects[group_label]:
            progress_done = int(project.get("segments", 0) or 0)
            progress_total = int(project.get("total_segments", 0) or 0)
            progress_percent = int(round((project.get("progress_ratio", 0.0) or 0.0) * 100))
            relative_time = (
                datetime.fromtimestamp(project["mtime"]).strftime("Hoy, %H:%M")
                if get_project_recency_group(project.get("mtime")) == "Hoy"
                else datetime.fromtimestamp(project["mtime"]).strftime("%d/%m/%Y")
                if project.get("mtime")
                else "Sin fecha"
            )
            card_col, action_col = st.columns([4.8, 1.25])
            with card_col:
                st.markdown(
                    f"""
                    <div class="dialog-project-card">
                      <strong>{html.escape(project.get("title", "Proyecto sin título"))}</strong>
                      <span>{int(project.get("media_count", 0) or 0)} material(es) · {html.escape(relative_time)} · {progress_percent}% transcrito</span>
                      <div class="dialog-progress-track">
                        <div class="dialog-progress-fill" style="width:{max(0, min(100, progress_percent))}%"></div>
                      </div>
                    </div>
                    """,
                    unsafe_allow_html=True,
                )
            with action_col:
                if st.button("Abrir", key=f"dialog_open_project_{project['id']}", use_container_width=True, type="primary"):
                    st.session_state.show_project_library_modal = False
                    open_existing_project(project["id"])
                if st.button("Renombrar", key=f"dialog_prepare_rename_{project['id']}", use_container_width=True):
                    st.session_state.dialog_project_rename_id = project["id"]
                if st.button("Eliminar", key=f"dialog_prepare_delete_{project['id']}", use_container_width=True):
                    st.session_state.dialog_project_delete_id = project["id"]

            if st.session_state.get("dialog_project_rename_id") == project["id"]:
                rename_col, save_col = st.columns([4, 1])
                with rename_col:
                    rename_title = st.text_input(
                        "Nuevo nombre",
                        value=project.get("title", ""),
                        key=f"dialog_project_rename_input_{project['id']}",
                    )
                with save_col:
                    st.write("")
                    if st.button("Guardar", key=f"dialog_project_rename_save_{project['id']}", use_container_width=True):
                        if rename_project(project["id"], rename_title):
                            st.session_state.dialog_project_rename_id = ""
                            st.rerun()
                        st.warning("No pude renombrar el proyecto.")

            if st.session_state.get("dialog_project_delete_id") == project["id"]:
                delete_col, cancel_col = st.columns([2, 1])
                with delete_col:
                    st.warning(f"¿Eliminar definitivamente “{project.get('title', 'este proyecto')}”?")
                with cancel_col:
                    if st.button("Cancelar", key=f"dialog_project_delete_cancel_{project['id']}", use_container_width=True):
                        st.session_state.dialog_project_delete_id = ""
                        st.rerun()
                    if st.button("Confirmar", key=f"dialog_project_delete_confirm_{project['id']}", use_container_width=True):
                        delete_project_state(project["id"])
                        if st.session_state.get("current_project_id") == project["id"]:
                            clear_active_project_state()
                        st.session_state.show_project_library_modal = False
                        st.rerun()


@st.dialog("Materiales", width="small")
def render_material_library_dialog():
    st.session_state.show_material_library_modal = True
    project_id = st.session_state.get("current_project_id")
    if not project_id:
        st.info("Primero abre un proyecto.")
        if st.button("Cerrar", key="dialog_close_material_library_without_project"):
            st.session_state.show_material_library_modal = False
            st.rerun()
        return

    project_meta = load_project_meta(project_id) or build_project_meta_from_session(project_id)
    materials = list_project_materials(project_id)
    active_material_id = st.session_state.get("active_material_id")
    show_add_material = bool(st.session_state.get("dialog_show_add_material_panel")) or not materials

    if not show_add_material:
        st.markdown('<div class="dialog-title dialog-title--highlight">Materiales del proyecto</div>', unsafe_allow_html=True)
        st.markdown(
            f'<div class="dialog-intro">Elige un material ya vinculado a <strong>{html.escape(project_meta.get("title", "este proyecto"))}</strong> para trabajarlo, o agrega uno nuevo.</div>',
            unsafe_allow_html=True,
        )
        active_material = next(
            (material for material in materials if material.get("id") == active_material_id),
            None,
        )
        ordered_materials = []
        if active_material:
            ordered_materials.append(active_material)
        ordered_materials.extend(
            material for material in materials
            if material.get("id") != active_material_id
        )
        other_materials = [material for material in ordered_materials if material.get("id") != active_material_id]

        if other_materials:
            st.markdown('<div class="dialog-chooser">', unsafe_allow_html=True)
            st.markdown('<div class="dialog-chooser-label">Cambiar material activo</div>', unsafe_allow_html=True)
            material_options = {
                f"{material.get('name', 'Material sin título')} · {format_size_mb(material.get('size_bytes', 0) or 0)}": material.get("id")
                for material in other_materials
            }
            selected_material_label = st.selectbox(
                "Elegir otro material",
                options=list(material_options.keys()),
                key="dialog_material_switcher",
                label_visibility="collapsed",
            )
            if st.button("Activar material elegido", key="dialog_activate_selected_material", use_container_width=True, type="primary"):
                selected_material_id = material_options.get(selected_material_label)
                if selected_material_id:
                    st.session_state.show_material_library_modal = False
                    activate_existing_material(selected_material_id)
                    st.rerun()
            st.markdown('</div>', unsafe_allow_html=True)

        for material in ordered_materials:
            progress_total = int(material.get("total_segments", 0) or 0)
            progress_done = int(material.get("transcribed_segments", 0) or 0)
            progress_percent = int(round((progress_done / progress_total) * 100)) if progress_total else 0
            is_active = material.get("id") == active_material_id
            card_class = "material-library-card material-library-card--active" if is_active else "material-library-card"
            action_label = "✓ Activo" if is_active else ""
            st.markdown(
                f"""
                <div class="{card_class}">
                  <div class="material-library-icon">A/V</div>
                  <div>
                    <strong>{html.escape(material.get("name", "Material sin título"))}</strong>
                    <span>{format_size_mb(material.get("size_bytes", 0) or 0)} · {progress_percent}% transcrito</span>
                    <div class="dialog-progress-track">
                      <div class="dialog-progress-fill" style="width:{max(0, min(100, progress_percent))}%"></div>
                    </div>
                  </div>
                  <div class="material-library-badge">{action_label}</div>
                </div>
                """,
                unsafe_allow_html=True,
            )

        if st.button("+ Agregar material nuevo", key="dialog_switch_to_add_material", use_container_width=True):
            st.session_state.dialog_show_add_material_panel = True
            st.rerun()
        return

    st.markdown('<div class="dialog-title dialog-title--highlight">Agregar material</div>', unsafe_allow_html=True)
    if materials:
        if st.button("‹ Volver a materiales del proyecto", key="dialog_back_to_material_list"):
            st.session_state.dialog_show_add_material_panel = False
            st.rerun()

    st.markdown(
        f'<div class="dialog-intro">Arrastra un archivo o haz clic. Quedará vinculado a <strong>{html.escape(project_meta.get("title", "este proyecto"))}</strong>.</div>',
        unsafe_allow_html=True,
    )

    st.markdown(
        f"""
        <div class="dialog-upload-drop">
          <strong>Suelta tu audio o video aquí</strong>
          <span>mp3, m4a, wav o mp4 · quedará vinculado a <strong>{html.escape(project_meta.get("title", "este proyecto"))}</strong></span>
        </div>
        """,
        unsafe_allow_html=True,
    )
    uploaded_file = st.file_uploader(
        "Selecciona un archivo",
        type=["mp3", "m4a", "wav", "mp4"],
        key=f"dialog_project_media_uploader_{st.session_state.media_uploader_version}",
        help="El archivo quedará vinculado al proyecto abierto y podrá quedar activo de inmediato.",
    )
    if not uploaded_file:
        return

    st.markdown(
        f"""
        <div class="dialog-file-card">
          <div class="material-library-icon">A/V</div>
          <div>
            <strong>{html.escape(uploaded_file.name)}</strong>
            <span>{format_size_mb(uploaded_file.size)} seleccionado</span>
          </div>
        </div>
        """,
        unsafe_allow_html=True,
    )

    if uploaded_file.size > WHISPER_FILE_LIMIT_BYTES:
        upload_max_mb, upload_use_silence = get_dialog_quick_split_settings("dialog_uploaded_material")
        estimated_segments = max(1, int((uploaded_file.size + (int(upload_max_mb) * BYTES_PER_MB) - 1) // (int(upload_max_mb) * BYTES_PER_MB)))
        render_dialog_split_recommendation(uploaded_file.size, upload_max_mb, upload_use_silence)
        with st.expander("Ajustar división", expanded=False):
            upload_max_mb, upload_use_silence = render_dialog_split_settings(
                uploaded_file.name,
                uploaded_file.size,
                "dialog_uploaded_material",
            )
            estimated_segments = max(1, int((uploaded_file.size + (int(upload_max_mb) * BYTES_PER_MB) - 1) // (int(upload_max_mb) * BYTES_PER_MB)))
        if st.button(f"Dividir en {estimated_segments} segmentos y activar", key="dialog_split_uploaded_material_btn", type="primary", use_container_width=True):
            with st.spinner("Guardando y dividiendo archivo grande..."):
                try:
                    activate_quick_split_segments(uploaded_file, upload_max_mb, upload_use_silence)
                    st.session_state.dialog_show_add_material_panel = False
                    st.session_state.show_material_library_modal = False
                    st.rerun()
                except Exception as exc:
                    st.error(f"No pude dividir el archivo: {exc}")
    else:
        if st.button("Activar este material", key="dialog_activate_uploaded_material_btn", type="primary", use_container_width=True):
            activate_uploaded_material(uploaded_file)
            st.session_state.dialog_show_add_material_panel = False
            st.session_state.show_material_library_modal = False
            st.rerun()

def main():
    init_session_state()

    if not check_ffmpeg_once():
        st.error(
            "No pude activar ffmpeg. Cierra la app y vuelve a abrirla para completar la preparación inicial."
        )
        st.stop()

    mode_choice = st.session_state.sidebar_model_choice
    installer_message = None
    if installer_flow_is_active(mode_choice):
        if st.session_state.current_project_id:
            installer_message = (
                "Instala primero el modelo local de Whisper para continuar con la transcripción de este proyecto. "
                "Tus audios seguirán procesándose localmente una vez terminada la descarga inicial."
            )
        else:
            installer_message = (
                "Antes de comenzar, instala el modelo local de transcripción que usará esta beta. "
                "Se descargará una sola vez y luego todo el procesamiento quedará en local."
            )

    render_app_header()
    highlight_transcribe_buttons()
    render_code_selection_from_editor(runtime_only=True)

    if installer_message:
        render_installer_gate_page(installer_message, mode_choice)
        return

    keep_sidebar_accessible()

    with st.sidebar:
        st.markdown('<div id="audioscript-sidebar-top-anchor"></div>', unsafe_allow_html=True)
        saved_projects = list_saved_projects()
        latest_project = get_latest_saved_project()
        current_project = next(
            (
                project
                for project in saved_projects
                if project["id"] == st.session_state.current_project_id
            ),
            None,
        )
        current_project_meta = (
            load_project_meta(st.session_state.current_project_id)
            if st.session_state.current_project_id
            else None
        )
        current_material_record = (
            get_active_material_record(current_project_meta)
            if current_project_meta
            else None
        )

        if st.session_state.current_project_id:
            if not st.session_state.active_audio_path:
                ready_title = "Falta agregar audio o video"
                ready_detail = (
                    "La sesión rápida ya está abierta; ahora agrega el material que vas a transcribir."
                    if st.session_state.get("quick_session_mode")
                    else "El proyecto está abierto; ahora agrega el material que vas a transcribir."
                )
                ready_class = " sidebar-ready-strip--warning"
            elif not is_project_preparation_confirmed(current_project_meta):
                ready_title = "Revisa la preparación"
                ready_detail = "Confirma modo, idioma, segmentos y modelo antes de iniciar o retomar la transcripción."
                ready_class = " sidebar-ready-strip--warning"
            else:
                ready_title = "Listo para transcribir"
                ready_detail = (
                    "Sesión rápida lista, material activo y configuración mínima disponible."
                    if st.session_state.get("quick_session_mode")
                    else "Proyecto abierto, material activo y configuración mínima disponible."
                )
                ready_class = ""

            st.markdown(
                f"""
                <div class="sidebar-ready-strip{ready_class}">
                  <div class="sidebar-ready-dot"></div>
                  <div>
                    <strong>{html.escape(ready_title)}</strong>
                    <span>{html.escape(ready_detail)}</span>
                  </div>
                </div>
                """,
                unsafe_allow_html=True,
            )

        render_sidebar_flow_steps(current_project_meta)

        render_sidebar_primary_action(get_contextual_sidebar_action(current_project_meta))
        if st.session_state.get("scroll_sidebar_to_top"):
            render_dom_script(
                """
                setTimeout(() => {
                  const doc = document;
                  const anchor = doc.querySelector('#audioscript-sidebar-top-anchor');
                  const sidebar = doc.querySelector('[data-testid="stSidebar"]');
                  const sidebarBody = doc.querySelector('[data-testid="stSidebarContent"]');
                  if (anchor) {
                    anchor.scrollIntoView({ behavior: 'smooth', block: 'start' });
                  }
                  if (sidebarBody) {
                    sidebarBody.scrollTop = 0;
                    sidebarBody.scrollTo({ top: 0, behavior: 'smooth' });
                  }
                  if (sidebar && sidebar !== sidebarBody) {
                    sidebar.scrollTop = 0;
                    sidebar.scrollTo({ top: 0, behavior: 'smooth' });
                  }
                }, 250);
                """,
            )
            st.session_state.scroll_sidebar_to_top = False
        if current_project:
            st.markdown('<div class="sidebar-section"><div class="sidebar-section-title">Proyecto</div>', unsafe_allow_html=True)
            project_options = {project["title"]: project["id"] for project in saved_projects} if saved_projects else {}
            ordered_titles = [project["title"] for project in saved_projects] if saved_projects else [current_project["title"]]
            current_title = current_project["title"]
            selected_title = st.selectbox(
                "Proyecto abierto",
                options=ordered_titles,
                index=ordered_titles.index(current_title) if current_title in ordered_titles else 0,
                key="sidebar_project_picker",
                label_visibility="collapsed",
            )
            selected_project_id = project_options.get(selected_title, current_project["id"])
            if selected_project_id != st.session_state.get("current_project_id"):
                open_existing_project(selected_project_id)
            st.markdown(
                f'<div class="sidebar-project-picker-note">{int(current_project.get("media_count", 0) or 0)} material(es) vinculados · {int(round((current_project.get("progress_ratio", 0.0) or 0.0) * 100))}% transcrito</div>',
                unsafe_allow_html=True,
            )
            quick_link_col1, quick_link_col2 = st.columns(2)
            with quick_link_col1:
                open_library = st.button("Ver biblioteca completa", key="open_project_library_link_btn", use_container_width=True)
            with quick_link_col2:
                quick_session = st.button("Sesión rápida", key="open_quick_session_btn", use_container_width=True)
            if open_library:
                st.session_state.show_project_library_modal = True
                st.rerun()
            if quick_session:
                create_quick_project()
                st.rerun()
            st.markdown("</div>", unsafe_allow_html=True)

        if st.session_state.current_project_id:
            st.markdown('<div class="sidebar-section"><div class="sidebar-section-title">Material para transcribir</div>', unsafe_allow_html=True)
            if current_material_record:
                st.markdown(
                    f"""
                    <div class="sidebar-material-card">
                      <div class="sidebar-material-icon">A/V</div>
                      <div>
                        <strong>{html.escape(current_material_record.get("name", st.session_state.active_audio_name))}</strong>
                        <span>{format_size_mb(current_material_record.get("size_bytes", 0) or 0)} · {format_duration(current_material_record.get("duration_seconds"))}</span>
                        <span>Material local vinculado al proyecto y listo para revisión o transcripción.</span>
                      </div>
                      <div class="sidebar-material-badge">Activo</div>
                    </div>
                    """,
                    unsafe_allow_html=True,
                )
            else:
                st.markdown(
                    """
                    <div class="sidebar-prep-card">
                      <strong>Aún no hay material activo.</strong>
                      <span>Sube o activa un audio/video desde aquí para convertirlo en el centro de trabajo del proyecto.</span>
                    </div>
                    """,
                    unsafe_allow_html=True,
                )
            material_button_label = (
                f"Elegir material ({current_project['media_count']})"
                if current_material_record and int(current_project.get("media_count", 0) or 0) > 1
                else "Cambiar o agregar material"
                if current_material_record
                else "Cargar audio o video"
            )
            if st.button(material_button_label, key="open_material_library_btn", use_container_width=True):
                st.session_state.show_material_library_modal = True
                st.session_state.dialog_show_add_material_panel = not bool(current_material_record)
                st.rerun()
            split_summary = st.session_state.get("quick_split_last_summary") or {}
            if split_summary:
                silence_adjustments = split_summary.get("silence_adjustments", 0)
                silence_label = (
                    f" {silence_adjustments} corte(s) ajustados a pausas naturales."
                    if split_summary.get("used_silence_detection")
                    else " Sin pausas detectadas cerca de los cortes."
                )
                st.success(
                    f"Quick Split: {split_summary.get('segment_count', 0)} segmento(s) creados desde {split_summary.get('source_name', 'archivo grande')}.{silence_label}"
                )
            st.markdown("</div>", unsafe_allow_html=True)

        st.markdown('<div class="sidebar-section"><div class="sidebar-section-title">Configuración</div>', unsafe_allow_html=True)
        render_preparation_summary(current_project_meta)

        st.markdown('<div id="audioscript-preparation-panel-anchor"></div>', unsafe_allow_html=True)
        if st.session_state.get("scroll_to_preparation_panel"):
            render_dom_script(
                """
                setTimeout(() => {
                  const doc = document;
                  const anchor = doc.querySelector('#audioscript-preparation-panel-anchor');
                  const sidebar = doc.querySelector('[data-testid="stSidebar"]');
                  if (anchor) {
                    anchor.scrollIntoView({ behavior: 'smooth', block: 'start' });
                  } else if (sidebar) {
                    sidebar.scrollTo({ top: sidebar.scrollHeight, behavior: 'smooth' });
                  }
                }, 250);
                """,
            )
            st.session_state.scroll_to_preparation_panel = False

        with st.expander("Preparación para la transcripción", expanded=bool(st.session_state.get("expand_preparation_panel"))):
            render_sidebar_minute_ring(st.session_state.sidebar_segment_mins)
            selected_processing_mode = render_segmented_or_fallback(
                "Modo de procesamiento",
                options=list(PROCESSING_MODES),
                current_value=normalize_processing_mode(st.session_state.sidebar_mode),
                key="prep_sidebar_mode",
            )
            st.session_state.sidebar_mode = normalize_processing_mode(
                selected_processing_mode or st.session_state.sidebar_mode
            )
            processing_mode_help = (
                "Solo para audios muy cortos."
                if st.session_state.sidebar_mode == MODE_COMPLETE
                else "Recomendado para entrevistas y análisis cualitativo."
            )
            st.markdown(
                f'<div class="sidebar-control-help">{processing_mode_help}</div>',
                unsafe_allow_html=True,
            )

            st.session_state.sidebar_language_choice = render_segmented_or_fallback(
                "Idioma",
                options=list(TRANSCRIPTION_LANGUAGES),
                current_value=(
                    st.session_state.sidebar_language_choice
                    if st.session_state.sidebar_language_choice in TRANSCRIPTION_LANGUAGES
                    else "Automático"
                    if st.session_state.sidebar_language_choice == "Detección automática"
                    else "Español"
                ),
                key="prep_sidebar_language",
            )

            segment_slider_col, segment_value_col = st.columns([4, 1])
            with segment_slider_col:
                st.session_state.sidebar_segment_mins = st.slider(
                    "Tamaño de segmentos",
                    min_value=1,
                    max_value=30,
                    value=int(st.session_state.sidebar_segment_mins),
                    key="prep_sidebar_segment_mins",
                )
            with segment_value_col:
                st.markdown(
                    f'<div class="sidebar-slider-value">{int(st.session_state.sidebar_segment_mins)} min</div>',
                    unsafe_allow_html=True,
                )

            st.session_state.sidebar_model_choice = render_segmented_or_fallback(
                "Modelo Whisper",
                options=["tiny", "base", "small", "medium", "large"],
                current_value=(
                    st.session_state.sidebar_model_choice
                    if st.session_state.sidebar_model_choice in ["tiny", "base", "small", "medium", "large"]
                    else DEFAULT_WHISPER_MODEL
                ),
                key="prep_sidebar_model",
                help_text="Modelos más grandes pueden ser más precisos, pero tardan más.",
            )
            st.markdown(
                '<div class="sidebar-control-help">Equilibrio recomendado para entrevistas.</div>',
                unsafe_allow_html=True,
            )
            if st.button("Confirmar preparación", key="confirm_preparation_btn", type="primary", use_container_width=True):
                confirm_project_preparation()
                st.session_state.expand_preparation_panel = False
                st.session_state.scroll_to_preparation_panel = False
                st.session_state.scroll_sidebar_to_top = True
                st.rerun()
        st.markdown("</div>", unsafe_allow_html=True)

        st.markdown('<div class="sidebar-section">', unsafe_allow_html=True)
        with st.expander("Lectura y formato", expanded=True):
            selected_transcription_format = render_segmented_or_fallback(
                "Formato de transcripción",
                options=["Limpia", "Verbatim"],
                current_value=(
                    st.session_state.sidebar_trans_type
                    if st.session_state.sidebar_trans_type in ["Limpia", "Verbatim"]
                    else "Limpia"
                    if "Limpia" in st.session_state.sidebar_trans_type
                    else "Verbatim"
                ),
                key="reading_sidebar_trans_type",
            )
            st.session_state.sidebar_trans_type = (
                selected_transcription_format
                or st.session_state.sidebar_trans_type
            )
            transcription_format_help = (
                "Literal con muletillas, pausas, risas."
                if st.session_state.sidebar_trans_type == "Verbatim"
                else "Sin muletillas ni repeticiones."
            )
            st.markdown(
                f'<div class="sidebar-control-help">{transcription_format_help}</div>',
                unsafe_allow_html=True,
            )
            font_slider_col, font_value_col = st.columns([4, 1])
            with font_slider_col:
                st.session_state.sidebar_editor_font_size = st.slider(
                    "Tamaño de letra del editor",
                    min_value=12,
                    max_value=28,
                    value=int(st.session_state.sidebar_editor_font_size),
                    key="reading_sidebar_font_size",
                )
            with font_value_col:
                st.markdown(
                    f'<div class="sidebar-slider-value">{int(st.session_state.sidebar_editor_font_size)} px</div>',
                    unsafe_allow_html=True,
                )
            st.text_area(
                "Nombres y términos técnicos",
                placeholder="Ej: Juan Pérez, fenomenología, ATLAS.ti...",
                height=100,
                key="main_custom_terms",
                help="Ayuda a Whisper a respetar nombres propios y vocabulario especializado.",
            )
            st.caption("Ayuda a Whisper a respetar nombres propios y vocabulario especializado.")

        with st.expander("Guía de trabajo", expanded=False):
            st.markdown(
                """
                **1. Abre o crea un proyecto**

                Un proyecto es el contenedor de trabajo. Puede reunir varios audios o videos, sus segmentos, memos,
                códigos y exportaciones. Antes de subir material, asegúrate de que el proyecto correcto esté abierto.

                **2. Selecciona y activa un material**

                Usa `Seleccionar y activar material` para subir un audio o video. El archivo queda vinculado al
                proyecto abierto. Si ya existe en el proyecto, puedes activarlo desde `Ver audios/videos del proyecto`.

                **3. Revisa la preparación**

                En `Preparación rápida` confirma el modo de trabajo, idioma, tamaño de segmentos y modelo Whisper.
                Para entrevistas largas o análisis cualitativo, usa modo `Segmentado`.

                **4. Archivos grandes**

                Si el archivo supera el límite recomendado, usa `Quick Split`. AudioScript lo divide en partes
                temporales y puede buscar silencios para evitar cortes a la mitad de una palabra.

                **5. Transcribe, lee y analiza**

                En la mesa central puedes transcribir, corregir y seleccionar frases. A la derecha está la
                `Mesa de análisis` para hacer memoing y pre-codificación mientras lees.

                **6. Exporta el avance**

                Puedes exportar aunque no hayas terminado todos los segmentos. El archivo Word incluye transcripción,
                memos y códigos guardados.
                """
            )

        with st.expander("Mantenimiento y limpieza", expanded=False):
            st.markdown(
                '<div class="sidebar-danger-note">Estas acciones afectan el proyecto abierto. Úsalas solo si quieres empezar desde cero.</div>',
                unsafe_allow_html=True,
            )
            confirm_cleanup = st.checkbox("Confirmo que quiero limpiar el proyecto activo")
            if st.button("Limpiar proyecto activo", disabled=not confirm_cleanup):
                delete_project_state(st.session_state.current_project_id)
                reset_transcription_state()
                st.session_state.current_project_id = None
                st.session_state.uploaded_file_id = None
                st.session_state.active_audio_name = ""
                st.session_state.active_audio_path = ""
                st.session_state.main_doc_title = "Transcripcion"
                st.session_state.project_description = ""
                st.session_state.quick_split_last_summary = {}
                st.session_state.show_new_project_form = False
                st.session_state.show_new_project_uploader = False
                st.session_state.media_uploader_version += 1
                st.session_state.autosave_last_saved = ""
                st.rerun()
        st.markdown("</div>", unsafe_allow_html=True)

    if st.session_state.show_project_library_modal:
        render_project_library_dialog()
    if st.session_state.show_material_library_modal:
        render_material_library_dialog()

    if st.session_state.get("scroll_to_transcription"):
        st.info("La mesa central de transcripción está lista más abajo en esta misma pantalla.")
        st.session_state.scroll_to_transcription = False

    mode = normalize_processing_mode(st.session_state.sidebar_mode)
    language_choice = st.session_state.sidebar_language_choice
    model_choice = st.session_state.sidebar_model_choice
    trans_type = st.session_state.sidebar_trans_type
    editor_font_size = st.session_state.sidebar_editor_font_size

    if not st.session_state.current_project_id:
        st.info("Usa el panel lateral para crear un proyecto o entrar a Modo rápido y comenzar.")
        render_workspace_watermarks()
        return

    if not st.session_state.active_audio_path or not os.path.exists(st.session_state.active_audio_path):
        if st.session_state.get("quick_session_mode"):
            st.info("Sesión rápida abierta. Agrega un audio o video desde el panel lateral para empezar la transcripción.")
        else:
            st.info("Proyecto abierto. Agrega un audio o video desde el panel lateral para empezar la transcripción.")
        render_workspace_watermarks()
        return

    apply_editor_styles(
        editor_font_size,
        reading_mode=st.session_state.get("reading_mode_enabled", False),
    )

    temp_audio_path = st.session_state.active_audio_path
    try:
        ensure_project_chunks(temp_audio_path)
    except RuntimeError as exc:
        st.session_state.chunks = []
        st.session_state.chunks_prepared = False
        st.error(str(exc))
        st.info(
            "Tu proyecto y sus transcripciones siguen intactos. Puedes volver a dividir el audio "
            "desde el botón que aparecerá en la mesa de transcripción."
        )
    memos_df = pd.DataFrame(st.session_state.memos)
    codes_df = build_codes_dataframe()

    project_context_complete = (
        st.session_state.main_doc_title.strip()
        and st.session_state.main_doc_title.strip() != "Transcripcion"
        and st.session_state.project_description.strip()
    )
    if not project_context_complete:
        st.markdown(
            """
            <div class="project-context-card">
              <div class="project-context-card__title">Ficha del proyecto y contexto</div>
              <div class="project-context-card__hint">Completa esta ficha antes de exportar: su título encabezará el documento Word.</div>
            </div>
            """,
            unsafe_allow_html=True,
        )
    with st.expander(
        "Ficha del proyecto y contexto",
        expanded=not bool(project_context_complete),
    ):
        meta_col, context_col = st.columns([1, 1.2])
        with meta_col:
            ensure_transcription_stamp()
            st.text_input(
                "Título de la transcripción",
                key="main_doc_title",
            )
            st.date_input(
                "Fecha del evento",
                value=st.session_state.main_event_date,
                key="main_event_date",
            )
            st.text_input(
                "Fecha de transcripción",
                key="main_transcription_date",
                disabled=True,
                help="AudioScript registra esta fecha automáticamente al iniciar la transcripción.",
            )
        with context_col:
            st.text_area(
                "Objetivo del evento y contexto",
                height=126,
                key="project_description",
                placeholder="",
            )

    render_project_context_card()

    doc_title = st.session_state.main_doc_title
    event_date = (
        st.session_state.main_event_date.strftime("%d/%m/%Y")
        if st.session_state.main_event_date
        else "No especificada"
    )
    transcription_date = st.session_state.main_transcription_date
    custom_terms = st.session_state.main_custom_terms
    segment_mins = st.session_state.sidebar_segment_mins

    st.caption(
        f"Guardado automático activo. Último guardado: {st.session_state.autosave_last_saved or 'pendiente'}"
    )

    if mode == MODE_COMPLETE:
        workspace_col, side_col = st.columns([3.6, 1.4], gap="large")
        render_keyboard_shortcuts(False)
        with workspace_col:
            open_card("Transcripción Completa", tone="transcription")
            render_transcript_toolbar_title("Documento completo")
            if st.button("Iniciar Transcripción Completa", type="primary"):
                with st.spinner("Transcribiendo audio completo..."):
                    try:
                        mark_transcription_started()
                        text = transcribe_audio(
                            temp_audio_path,
                            model_choice,
                            language_choice,
                            custom_terms,
                            trans_type,
                        )
                        st.session_state.last_transcription = text
                        save_project_state()
                    except Exception as exc:
                        st.error(f"Error durante la transcripción: {exc}")

            if st.session_state.last_transcription:
                toolbar_col1, toolbar_col2, toolbar_col3 = st.columns([1.4, 1.4, 0.35])
                with toolbar_col1:
                    complete_find_term = st.text_input(
                        "Buscar en transcripción",
                        key="complete_find_term",
                    )
                with toolbar_col2:
                    complete_replace_term = st.text_input(
                        "Reemplazar con",
                        key="complete_replace_term",
                    )
                with toolbar_col3:
                    if st.button("↳", key="complete_replace_btn", help="Aplicar reemplazo"):
                        if complete_find_term:
                            st.session_state.last_transcription = (
                                st.session_state.last_transcription.replace(
                                    complete_find_term,
                                    complete_replace_term,
                                )
                            )
                            save_project_state()
                            st.rerun()

                render_transcript_panel_label(
                    "Panel único de transcripción",
                    "Lee, corrige y selecciona frases en este mismo cuadro antes de exportar o codificar.",
                )
                reading_col, reading_hint_col = st.columns([0.9, 2.1])
                with reading_col:
                    render_reading_mode_toggle("complete_reading_mode_toggle")
                with reading_hint_col:
                    render_reading_mode_hint()
                final_text = st.text_area(
                    "Edite el texto si es necesario",
                    value=st.session_state.last_transcription,
                    height=483,
                )
                st.session_state.last_transcription = final_text
                render_code_selection_from_editor("Edite el texto si es necesario")
                render_downloads(
                    final_text,
                    doc_title,
                    event_date,
                    transcription_date,
                )
            else:
                st.info("Cuando inicies la transcripción completa, el texto aparecerá aquí.")
            close_card()

        with side_col:
            open_analysis_desk("documento completo")
            render_code_manager(0)
            render_memo_manager(0)
            close_analysis_desk()
        render_footer_bar()
        return

    if not st.session_state.chunks:
        if st.button("Dividir audio en segmentos", type="primary"):
            with st.spinner("Dividiendo audio..."):
                try:
                    rebuilt_chunks = rebuild_active_audio_segments(
                        temp_audio_path,
                        segment_mins,
                        st.session_state.current_project_id,
                        st.session_state.get("active_material_id"),
                    )
                    st.success(
                        f"Audio dividido en {len(rebuilt_chunks)} segmentos."
                    )
                    st.rerun()
                except Exception as exc:
                    st.error(f"Error al dividir el audio: {exc}")

    chunks = st.session_state.chunks
    ensure_segment_text_store(len(chunks))
    idx = st.session_state.current_chunk_idx

    if not chunks:
        st.info("Haz clic en 'Dividir audio en segmentos' para comenzar.")
        return

    if idx < len(chunks):
        pending_segment_action = st.session_state.get("pending_segment_action")
        if pending_segment_action == "advance":
            if st.session_state.current_segment_text.strip():
                sync_current_segment_text()
                next_idx = st.session_state.current_chunk_idx + 1
                if next_idx < len(chunks):
                    load_segment_text(next_idx)
                else:
                    st.session_state.current_chunk_idx = len(chunks)
                    st.session_state.current_segment_text = ""
            st.session_state.pending_segment_action = None
            save_project_state()
            chunks = st.session_state.chunks
            idx = st.session_state.current_chunk_idx
            if idx >= len(chunks):
                st.rerun()
        elif pending_segment_action == "reset":
            stored = ensure_segment_text_store(len(chunks))
            if 0 <= st.session_state.current_chunk_idx < len(stored):
                stored[st.session_state.current_chunk_idx] = ""
                st.session_state.segment_texts = stored
                st.session_state.transcript_segments = [text for text in stored if text.strip()]
            st.session_state.current_segment_text = ""
            st.session_state.pending_segment_action = None
            save_project_state()

        st.markdown(f"### Segmento {idx + 1} de {len(chunks)}")
        st.progress(idx / len(chunks))
        render_segment_navigation(chunks)
        if (
            st.session_state.chunk_segment_mins
            and st.session_state.chunk_segment_mins != segment_mins
        ):
            warning_col, action_col = st.columns([4.2, 1.1])
            with warning_col:
                st.info(
                    "Cambiaste los minutos por segmento. Para aplicar ese cambio, vuelve a dividir el audio."
                )
            with action_col:
                if st.button("Redividir", key="rebuild_segments_btn", type="primary", use_container_width=True):
                    with st.spinner("Volviendo a dividir audio..."):
                        try:
                            rebuilt_chunks = rebuild_active_audio_segments(
                                temp_audio_path,
                                segment_mins,
                                st.session_state.current_project_id,
                                st.session_state.get("active_material_id"),
                            )
                            st.success(f"Audio redividido en {len(rebuilt_chunks)} segmentos.")
                            st.rerun()
                        except Exception as exc:
                            st.error(f"No pude volver a dividir el audio: {exc}")
        render_keyboard_shortcuts(True)

        open_card("Control de Reproducción", tone="playback")
        render_audio_console_header(f"Segmento {idx + 1} de {len(chunks)}")
        render_segment_audio_player(chunks[idx])
        st.caption(
            "Atajos: `Ctrl+Shift+T` transcribe, `Ctrl+Enter` confirma, `Ctrl+Shift+R` reinicia, `Alt + ←/→` mueve el audio."
        )
        close_card()

        col_main, col_side = st.columns([3.6, 1.4], gap="large")

        with col_main:
            open_card("Transcripción", tone="transcription")
            render_transcript_toolbar_title(f"Segmento {idx + 1} · {doc_title}")
            toolbar_col1, toolbar_col2, toolbar_col3, toolbar_col4, toolbar_col5 = st.columns([1.3, 1.3, 0.32, 1.15, 1])
            with toolbar_col1:
                find_term = st.text_input("Buscar en transcripción", key=f"find_term_{idx}")
            with toolbar_col2:
                replace_term = st.text_input("Reemplazar con", key=f"replace_term_{idx}")
            with toolbar_col3:
                if st.button("↳", key=f"replace_btn_{idx}", help="Aplicar reemplazo"):
                    if find_term:
                        st.session_state.current_segment_text = (
                            st.session_state.current_segment_text.replace(find_term, replace_term)
                        )
                        save_project_state()
                        st.rerun()
            with toolbar_col4:
                if st.button("Confirmar y siguiente", key=f"confirm_next_btn_{idx}"):
                    if st.session_state.current_segment_text.strip():
                        st.session_state.pending_segment_action = "advance"
                        st.rerun()
                    st.warning("No hay texto para guardar.")
            with toolbar_col5:
                if st.button("Transcribir", key=f"transcribe_btn_{idx}", type="primary"):
                    with st.spinner("Whisper está transcribiendo..."):
                        try:
                            mark_transcription_started()
                            text = transcribe_audio(
                                chunks[idx],
                                model_choice,
                                language_choice,
                                custom_terms,
                                trans_type,
                            )
                            st.session_state.current_segment_text = text
                            save_project_state()
                            st.rerun()
                        except Exception as exc:
                            st.error(f"Error en transcripción: {exc}")

            if find_term:
                occurrences = st.session_state.current_segment_text.lower().count(find_term.lower())
                st.caption(f"Coincidencias encontradas en el fragmento: {occurrences}")

            render_transcript_panel_label(
                "Panel único de transcripción",
                "Lee, corrige y selecciona frases en este mismo cuadro antes de confirmar el segmento.",
            )
            reading_col, reading_hint_col = st.columns([0.9, 2.1])
            with reading_col:
                render_reading_mode_toggle(f"segment_reading_mode_toggle_{idx}")
            with reading_hint_col:
                render_reading_mode_hint()
            st.text_area(
                "Edite este fragmento antes de continuar",
                height=322,
                key="current_segment_text",
            )
            render_code_selection_from_editor("Edite este fragmento antes de continuar")

            sync_current_segment_text()
            partial_segments = get_ordered_segment_texts()
            partial_text = "\n\n".join(partial_segments)
            if partial_text.strip():
                st.markdown("#### Exportación parcial")
                st.caption(
                    "Puedes descargar el avance aunque todavía falten segmentos por transcribir."
                )
                render_downloads(
                    partial_text,
                    doc_title,
                    event_date,
                    transcription_date,
                    memos_df,
                    codes_df,
                )

            c1, c2 = st.columns([1, 1])
            with c1:
                if st.button("Reiniciar segmento"):
                    st.session_state.pending_segment_action = "reset"
                    st.rerun()
            with c2:
                st.metric("Palabras", len(st.session_state.current_segment_text.split()))

            stats = [
                ("Segmentos", f"{len(partial_segments)}/{len(chunks)}"),
                ("Memos", str(len(st.session_state.memos))),
                ("Códigos", str(len(st.session_state.codes))),
                ("Caracteres", str(len(st.session_state.current_segment_text))),
            ]
            st.markdown(
                '<div class="mini-stat-grid">' +
                "".join(
                    [
                        f'<div class="mini-stat"><div class="mini-stat__label">{label}</div><div class="mini-stat__value">{value}</div></div>'
                        for label, value in stats
                    ]
                ) +
                '</div>',
                unsafe_allow_html=True,
            )
            close_card()

        with col_side:
            open_analysis_desk(f"segmento {idx + 1}")
            render_code_manager(idx + 1)
            render_memo_manager(idx + 1)
            close_analysis_desk()
        render_footer_bar()
        return

    st.success("Transcripción completada")
    workspace_col, side_col = st.columns([3.6, 1.4], gap="large")
    full_text = "\n\n".join(get_ordered_segment_texts())
    render_keyboard_shortcuts(True)
    with workspace_col:
        open_card("Transcripción Consolidada", tone="transcription")
        render_transcript_toolbar_title("Documento consolidado")
        toolbar_col1, toolbar_col2, toolbar_col3 = st.columns([1.4, 1.4, 0.35])
        with toolbar_col1:
            final_find_term = st.text_input(
                "Buscar en transcripción",
                key="final_find_term",
            )
        with toolbar_col2:
            final_replace_term = st.text_input(
                "Reemplazar con",
                key="final_replace_term",
            )
        with toolbar_col3:
            if st.button("↳", key="final_replace_btn", help="Aplicar reemplazo"):
                if final_find_term:
                    updated_segments = [
                        segment.replace(final_find_term, final_replace_term)
                        for segment in get_ordered_segment_texts(include_blanks=True)
                    ]
                    st.session_state.segment_texts = updated_segments
                    st.session_state.transcript_segments = [
                        segment for segment in updated_segments if segment.strip()
                    ]
                    save_project_state()
                    st.rerun()

        render_transcript_panel_label(
            "Panel único de transcripción consolidada",
            "Lee, ajusta y selecciona frases en este mismo cuadro antes de descargar o codificar.",
        )
        reading_col, reading_hint_col = st.columns([0.9, 2.1])
        with reading_col:
            render_reading_mode_toggle("final_reading_mode_toggle")
        with reading_hint_col:
            render_reading_mode_hint()
        final_view = st.text_area(
            "Contenido consolidado",
            value=full_text,
            height=420,
        )
        render_code_selection_from_editor("Contenido consolidado")
        render_downloads(
            final_view,
            doc_title,
            event_date,
            transcription_date,
        )

        if st.button("Empezar de nuevo"):
            delete_project_state(st.session_state.current_project_id)
            reset_transcription_state()
            st.session_state.current_project_id = None
            st.session_state.uploaded_file_id = None
            st.session_state.active_audio_name = ""
            st.session_state.active_audio_path = ""
            st.session_state.main_doc_title = "Transcripcion"
            st.session_state.project_description = ""
            st.session_state.quick_split_last_summary = {}
            st.session_state.show_new_project_form = False
            st.session_state.show_new_project_uploader = False
            st.session_state.media_uploader_version += 1
            st.session_state.autosave_last_saved = ""
            st.rerun()
        close_card()

    with side_col:
        open_analysis_desk("transcripción consolidada")
        render_code_manager(0)
        render_memo_manager(0)
        close_analysis_desk()
    render_footer_bar()


if __name__ == "__main__":
    main()
